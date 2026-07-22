"""File front door: /api/studio/ingest-file — PDF/DOCX/txt/md → text → the shared IE→ER→graph pipeline.

Format handling is tested with REAL in-memory documents (python-docx builds a docx; pypdf builds a
blank pdf), not fixtures — so a parser bump that changes behavior fails here, not in production.
"""
import base64
import io

from fastapi.testclient import TestClient

import lattice_studio.server as srv
from lattice_studio.server import _doc_to_text, app

client = TestClient(app)


def _b64(data: bytes) -> str:
    return base64.b64encode(data).decode()


async def _fake_req(client_, method, url, json=None):
    # IE/ER/hellgraph all "reachable": IE returns one entity so the pipeline has something to write.
    if "ie-engine" in url:
        return {"entities": [{"text": "Guzman & Gomez", "type": "Organization"}], "relations": [], "claims": []}, None
    if "/resolve" in url:
        return {"replay_key": "er:x", "merged": 0, "review_queue": [],
                "entities": [{"entity_id": "ent:m0", "members": ["m0"], "size": 1,
                              "canonical": {"survivor": "m0", "name": "Guzman & Gomez"}}]}, None
    return {"ok": True}, None


def test_ingest_file_write_gate(monkeypatch):
    monkeypatch.setattr("lattice_studio.server.STUDIO_WRITE_TOKEN", "")
    r = client.post("/api/studio/ingest-file",
                    json={"filename": "a.txt", "content_b64": _b64(b"hello")})
    assert r.status_code == 503  # fail-closed


def test_ingest_file_rejects_bad_b64_empty_and_oversize(monkeypatch):
    monkeypatch.setattr("lattice_studio.server.STUDIO_WRITE_TOKEN", "secret")
    hdr = {"authorization": "Bearer secret"}
    assert client.post("/api/studio/ingest-file", headers=hdr,
                       json={"filename": "a.txt", "content_b64": "not!!base64??"}).status_code == 400
    assert client.post("/api/studio/ingest-file", headers=hdr,
                       json={"filename": "a.txt", "content_b64": ""}).status_code == 400
    monkeypatch.setattr("lattice_studio.server.MAX_FILE_BYTES", 4)
    assert client.post("/api/studio/ingest-file", headers=hdr,
                       json={"filename": "a.txt", "content_b64": _b64(b"12345")}).status_code == 413


def test_ingest_file_unsupported_type_names_the_alternative(monkeypatch):
    monkeypatch.setattr("lattice_studio.server.STUDIO_WRITE_TOKEN", "secret")
    r = client.post("/api/studio/ingest-file", headers={"authorization": "Bearer secret"},
                    json={"filename": "data.csv", "content_b64": _b64(b"a,b\n1,2")})
    assert r.status_code == 415
    assert "/api/studio/ingest" in r.json()["detail"]  # tabular data has its own door


def test_ingest_file_txt_runs_pipeline_with_file_provenance(monkeypatch):
    monkeypatch.setattr("lattice_studio.server.STUDIO_WRITE_TOKEN", "secret")
    monkeypatch.setattr(srv, "_req", _fake_req)
    raw = b"Guzman & Gomez opened 12 stores."
    r = client.post("/api/studio/ingest-file", headers={"authorization": "Bearer secret"},
                    json={"project": "team-x", "filename": "gyg.txt", "content_b64": _b64(raw)})
    assert r.status_code == 200
    b = r.json()
    assert b["projectCollection"] == "proj-teamx"
    assert b["written"]["nodes"] == 1
    # double provenance: file_sha (dropped bytes) AND doc_sha (extracted text) — both traceable
    import hashlib
    assert b["provenance"]["file_sha"] == hashlib.sha256(raw).hexdigest()
    assert b["provenance"]["doc_sha"] == hashlib.sha256(raw.decode().strip().encode()).hexdigest()
    assert b["provenance"]["filename"] == "gyg.txt"


def test_docx_conversion_includes_paragraphs_and_tables():
    from docx import Document
    d = Document()
    d.add_paragraph("Guzman & Gomez quarterly report.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Revenue"
    t.rows[0].cells[1].text = "$12m"
    buf = io.BytesIO()
    d.save(buf)
    text = _doc_to_text("report.docx", buf.getvalue())
    assert "Guzman & Gomez quarterly report." in text
    assert "Revenue | $12m" in text  # tables carry the numbers — must survive conversion


def test_blank_pdf_yields_422_not_500(monkeypatch):
    from pypdf import PdfWriter
    w = PdfWriter()
    w.add_blank_page(width=612, height=792)  # a page with no extractable text (scan-like)
    buf = io.BytesIO()
    w.write(buf)
    monkeypatch.setattr("lattice_studio.server.STUDIO_WRITE_TOKEN", "secret")
    r = client.post("/api/studio/ingest-file", headers={"authorization": "Bearer secret"},
                    json={"filename": "scan.pdf", "content_b64": _b64(buf.getvalue())})
    assert r.status_code == 422
    assert "OCR" in r.json()["detail"]


def test_corrupt_pdf_yields_422_not_500(monkeypatch):
    monkeypatch.setattr("lattice_studio.server.STUDIO_WRITE_TOKEN", "secret")
    r = client.post("/api/studio/ingest-file", headers={"authorization": "Bearer secret"},
                    json={"filename": "broken.pdf", "content_b64": _b64(b"%PDF-1.4 garbage")})
    assert r.status_code == 422
