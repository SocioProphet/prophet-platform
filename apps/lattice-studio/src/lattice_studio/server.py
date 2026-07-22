"""server.py — the Lattice Studio BFF.

Wraps lattice-studio's existing ``product_spine`` (pure functions) into a real HTTP service and AGGREGATES LIVE
data from the running fabric, project-scoped to Noetica ``proj-`` collections. This is the "make it live" service.

Design (production-grade):
  * CONCURRENT fan-out — all upstream calls run under a single ``asyncio.gather`` (not sequential awaits), so
    ``/api/studio`` latency is max(upstream) not sum(upstream).
  * GRACEFUL — every upstream is wrapped; a down/slow service degrades only its own section, never the response.
  * HONEST — the ``live`` map reports exactly which upstreams answered; ``degraded`` carries the reason.

Upstreams (real, in-cluster):
  * hellgraph-service  GET  /api/graph/stats          → Graph section (live stats)
  * search-orchestrator POST /v0/search/query          → Extraction/Sherlock (live federated results)
  * tritfabric         GET  /v1/registry              → Model catalog (live) — degrades until tritfabric is deployed
"""
from __future__ import annotations

import asyncio
import base64
import csv
import hashlib
import io
import json
import os
import re
import uuid
from datetime import datetime, timezone
from typing import Any

import httpx
import jwt
from fastapi import Depends, FastAPI, Header, HTTPException, Response
from pydantic import BaseModel

from lattice_studio import gaia, hdt, ontology, product_spine, shacl

SERVICE_VERSION = "0.2.0"
HELLGRAPH_URL = os.getenv("HELLGRAPH_URL", "http://hellgraph-service:8090")
TRITFABRIC_URL = os.getenv("TRITFABRIC_URL", "http://tritfabric:8750")
SEARCH_ORCH_URL = os.getenv("SEARCH_ORCH_URL", "http://search-orchestrator:8080")
# Document-ingestion pipeline upstreams: real NER/relation extraction + entity resolution.
IE_ENGINE_URL = os.getenv("IE_ENGINE_URL", "http://ie-engine:8080")
ER_URL = os.getenv("ER_URL", "http://entity-resolution:8080")
TIMEOUT = float(os.getenv("STUDIO_TIMEOUT", "5"))
# Write gate for /api/studio/extract (mutates the graph). Fail-closed: unset → writes refused.
STUDIO_WRITE_TOKEN = os.getenv("STUDIO_WRITE_TOKEN", "")
# Read gate tied to the SOVEREIGN identity plane: the HS256 secret socbase (GoTrue) signs its JWTs with.
# OPT-IN — unset → reads stay open (backward compatible); set → every read requires a valid socbase-issued
# bearer token. Governance without a bolt-on: Studio reads are gated by the same identity that runs the estate.
STUDIO_JWT_SECRET = os.getenv("STUDIO_JWT_SECRET", "")
# Evidence fabric — verified-compute RECEIPTS. Studio surfaces the replayable proof-of-work behind its services:
# not "the job ran" (a TEE attestation), but a sealed record of exactly WHAT ran and its verdict, per correlation.
EVIDENCE_RECEIPTS_URL = os.getenv("EVIDENCE_RECEIPTS_URL", "http://evidence-receipts:8080")
# The sovereign Spark execution backend (apps/spark-runner). When backend='spark' + entitled, execute dispatches
# the job here and chains its receipt. Unset/unreachable → the run degrades to the governed ledger (dispatched).
SPARK_RUNNER_URL = os.getenv("SPARK_RUNNER_URL", "")
SPARK_RUNNER_TOKEN = os.getenv("SPARK_RUNNER_TOKEN", "")
RECEIPT_SERVICES = [s.strip() for s in os.getenv(
    "STUDIO_RECEIPT_SERVICES",
    "hellgraph-service,lattice-studio,search-orchestrator,owl-reasoner,entity-resolution,eval-fabric-api",
).split(",") if s.strip()]
# Notebook runtime (lattice-forge, in the isolated sovereign-runtime namespace).
FORGE_URL = os.getenv("FORGE_URL", "http://lattice-forge.sovereign-runtime.svc.cluster.local:8870")
FORGE_TOKEN = os.getenv("FORGE_TOKEN", "")
FORGE_TIMEOUT = float(os.getenv("FORGE_TIMEOUT", "90"))
# Universal Compute Plane gateway (any compute, one governed door).
COMPUTE_GATEWAY_URL = os.getenv("COMPUTE_GATEWAY_URL", "http://compute-gateway:8080")
GATEWAY_TOKEN = os.getenv("GATEWAY_TOKEN", "")
GATEWAY_TIMEOUT = float(os.getenv("GATEWAY_TIMEOUT", "120"))

app = FastAPI(title="Lattice Studio BFF", version=SERVICE_VERSION)


def require_read(authorization: str = Header(default="")) -> dict[str, Any] | None:
    """FastAPI dependency for READ endpoints. If STUDIO_JWT_SECRET is unset, returns None (reads open). When set,
    verifies the Bearer JWT (HS256) with that secret and returns its claims (the caller identity); a missing,
    invalid, or expired token is a 401. Reads are thereby gated by sovereign identity, not an ad-hoc password."""
    if not STUDIO_JWT_SECRET:
        return None
    token = authorization.removeprefix("Bearer ").strip()
    if not token:
        raise HTTPException(status_code=401, detail="read requires a bearer token (STUDIO_JWT_SECRET is set)")
    try:
        return jwt.decode(token, STUDIO_JWT_SECRET, algorithms=["HS256"], options={"verify_aud": False})
    except jwt.PyJWTError as exc:
        raise HTTPException(status_code=401, detail=f"invalid token: {exc}") from exc


def proj_collection(project: str) -> str:
    """Mirror Noetica projectCollectionId — proj-<12 hex of the project id, dashes stripped>."""
    return "proj-" + re.sub(r"-", "", project)[:12]


def _first(d: dict[str, Any], *keys: str, default: str = "—") -> Any:
    for k in keys:
        if d.get(k):
            return d[k]
    return default


async def _req(client: httpx.AsyncClient, method: str, url: str, json: Any | None = None) -> tuple[Any, str | None]:
    """One resilient upstream call. Returns (json_or_None, error_or_None). Never raises."""
    try:
        r = await client.request(method, url, json=json)
        if r.status_code == 200:
            return r.json(), None
        return None, f"HTTP {r.status_code}"
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


def _sherlock_request(project: str) -> dict[str, Any]:
    """A schema-valid sherlock_search_request (see schemas/search/sherlock_search_request.schema.json)."""
    return {
        "query_id": f"studio-{proj_collection(project)}",
        "actor_id": "lattice-studio-bff",
        "text": project,
        "mode": "HYBRID",
        "scope": {"local_desktop": False, "cloud_workspace": True, "memory": True},
        "limit": 10,
    }


@app.get("/healthz")
def healthz() -> dict[str, Any]:
    return {"ok": True, "service": "lattice-studio", "version": SERVICE_VERSION}


@app.get("/api/studio")
async def studio(project: str = "default", _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    coll = proj_collection(project)
    spine = product_spine.demo_product_spine()  # the real integration object model

    # ── CONCURRENT live fan-out to the running fabric ──
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        (gstats, gerr), (reg, rerr), (sher, serr), (subg, _suberr), (rec, _recerr) = await asyncio.gather(
            _req(client, "GET", f"{HELLGRAPH_URL}/api/graph/stats"),
            _req(client, "GET", f"{TRITFABRIC_URL}/v1/registry"),
            _req(client, "POST", f"{SEARCH_ORCH_URL}/v0/search/query", json=_sherlock_request(project)),
            _req(client, "GET", f"{HELLGRAPH_URL}/api/graph/subgraph?label={coll}&limit=300"),
            _req(client, "GET", f"{EVIDENCE_RECEIPTS_URL}/v1/receipts/recent?service=hellgraph-service&limit=10"),
        )
    degraded = {k: v for k, v in {"graph": gerr, "models": rerr, "extraction": serr}.items() if v}

    # ── The MOAT header, computed live on every load: epistemic distribution + provenance coverage across the
    # project's facts, plus whether the verified-compute evidence fabric is answering. This governance readout
    # rides ABOVE every section — the proof-carrying identity of the whole workspace, not just the graph panel.
    epi_nodes = [_map_node(n) for n in ((subg.get("nodes") if isinstance(subg, dict) else None) or [])]
    epi_dist: dict[str, int] = {}
    prov_have = 0
    for n in epi_nodes:
        epi_dist[n["epistemic_mode"]] = epi_dist.get(n["epistemic_mode"], 0) + 1
        if n.get("source"):
            prov_have += 1
    moat = {
        "epistemic_distribution": epi_dist,
        "fact_count": len(epi_nodes),
        "provenance_coverage": round(prov_have / len(epi_nodes), 3) if epi_nodes else 0.0,
        "verified_compute": rec is not None,
        "receipts_recent": len((rec.get("items") if isinstance(rec, dict) else None) or []),
        "governed_writes": bool(STUDIO_WRITE_TOKEN),
        "read_auth": bool(STUDIO_JWT_SECRET),
    }

    # ── Workbench: product_spine object model, bound to the Noetica project ──
    session = dict(spine.get("notebookSession", {}))
    session["projectId"] = project
    notebooks = [{
        "id": _first(session, "notebookSessionId", "id", default="nb-session"),
        "name": _first(session, "name", "title", default="Notebook session"),
        "runtime": _first(spine.get("runtimeAsset", {}), "runtimeAssetId", "name", default="prophet-python-ml"),
        "kernel": "python3", "status": "idle", "updatedAt": _first(session, "createdAt", default=""),
        "cells": 0, "collaborators": ["you"], "projectCollection": coll,
    }]
    data = [
        {"id": _first(d, "catalogAssetId", "id"), "name": _first(d, "name", "title"), "kind": "dataset",
         "catalog": "prophet-core-catalog", "governed": True,
         "lineage": [_first(d, "reproduceCommand", "reproduce_command", default="lineage")]}
        for d in (spine.get("catalogAsset", {}), spine.get("dataProduct", {})) if d
    ]
    models: list[dict[str, Any]] = []
    if reg:
        arts = reg.get("artifacts", reg) if isinstance(reg, dict) else reg
        for a in (arts if isinstance(arts, list) else []):
            models.append({"id": _first(a, "id", "artifact_id", default="m"), "name": _first(a, "name", "id"),
                           "task": _first(a, "task", default="model"), "stage": _first(a, "stage", "state", default="candidate"),
                           "servable": True, "metrics": []})
    if not models:
        for m in (spine.get("factsheet", {}), spine.get("promotionCandidate", {})):
            if m:
                models.append({"id": _first(m, "id", "candidateId"), "name": _first(m, "id", "candidateId"),
                               "task": "governed", "stage": "staged", "lineage": m.get("lineageRefs", [])})
    experiments = [
        {"id": _first(e, "id"), "title": _first(e, "title", "id"), "reproducible": True,
         "provenance": "in-toto + lockfile", "createdAt": "", "rerunnable": True}
        for e in (spine.get("evaluationBundle", {}), spine.get("publicationArtifact", {})) if e
    ]

    # ── Graph: LIVE hellgraph stats ──
    graph: list[dict[str, Any]] = []
    if isinstance(gstats, dict):
        for k, v in gstats.items():
            if isinstance(v, (int, float, str)):
                graph.append({"id": f"g-{k}", "label": k, "value": v})
    graph.append({"id": "g-engine", "label": "Engine", "value": "HellGraph", "hint": f"live · {HELLGRAPH_URL}"})

    # ── Knowledge engineering — LIVE where a service exists, honest lib/spec status otherwise ──
    sher_hits = len(sher.get("results", sher) if isinstance(sher, dict) else (sher or [])) if sher else 0
    extraction = [
        {"id": "x-holmes", "name": "Holmes — entities & relations", "engine": "holmes",
         "kind": "claim reasoning (Propose→Explain→Verify)", "status": "idle", "target": coll},
        {"id": "x-sherlock", "name": "Sherlock — federated search", "engine": "sherlock", "kind": "federated retrieval",
         "status": ("done" if sher is not None else "idle"), "extracted": sher_hits, "target": coll},
    ]
    ontology = [
        {"id": "o-1", "name": "Ontogenesis modules (RDF/OWL + SHACL)", "kind": "class", "engine": "ontogenesis", "aligned": True},
        {"id": "o-epi", "name": "Epistemology (epi.ttl / epistemic_mode)", "kind": "axiom", "engine": "ontogenesis", "aligned": True},
    ]
    retrieval = [
        {"id": "r-fiber", "name": "Fibered retrieval (PageIndex ⊕ HellGraph)", "method": "fiber", "engine": "fibered-retrieval", "scope": "project", "ready": True},
        {"id": "r-grag", "name": "Graph-RAG", "method": "graph-rag", "engine": "hellgraph", "scope": "project", "ready": gstats is not None},
        {"id": "r-topic", "name": "Topic packs (/topic)", "method": "topic", "engine": "slash-topics", "scope": "project", "ready": True},
        {"id": "r-sem", "name": "Semantic + lexical (sheaf)", "method": "vector", "engine": "noetica", "scope": "chat + project", "ready": True},
    ]
    generation = [
        {"id": "gn-1", "name": "New-Hope grounded synthesis", "engine": "new-hope", "kind": "synthesis", "status": "idle", "grounded": True},
    ]

    return {
        "project": project, "projectCollection": coll,
        "moat": moat,
        "notebooks": notebooks, "data": data, "models": models, "tuning": [], "experiments": experiments,
        "extraction": extraction, "ontology": ontology, "graph": graph, "retrieval": retrieval, "generation": generation,
        "live": {"hellgraph": gstats is not None, "tritfabric": reg is not None,
                 "search_orchestrator": sher is not None, "evidence": rec is not None},
        "degraded": (degraded or None),
    }


def _receipt_verdict(item: dict[str, Any]) -> Any:
    rc = item.get("receipt") if isinstance(item.get("receipt"), dict) else {}
    return item.get("verdict") or rc.get("verdict")


@app.get("/api/studio/receipts")
async def receipts(limit: int = 12, _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Verified-compute RECEIPTS from the evidence fabric — the replayable proof-of-work behind Studio's services,
    aggregated across the estate. Each is a sealed record (fetch the full bundle at /v1/receipts/{service}/{cid}):
    not a TEE's 'it ran privately', but 'here is exactly WHAT ran and its verdict'. No incumbent studio has this.
    Graceful: a down evidence fabric or service degrades only its own slice."""
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        results = await asyncio.gather(*[
            _req(client, "GET", f"{EVIDENCE_RECEIPTS_URL}/v1/receipts/recent?service={s}&limit={limit}")
            for s in RECEIPT_SERVICES
        ])
    out: list[dict[str, Any]] = []
    live: dict[str, bool] = {}
    for svc, (res, err) in zip(RECEIPT_SERVICES, results):
        live[svc] = err is None
        items = (res.get("items") if isinstance(res, dict) else None) or []
        for it in (items if isinstance(items, list) else [])[:limit]:
            it = it if isinstance(it, dict) else {}
            cid = str(it.get("correlation_id") or it.get("id") or "")
            out.append({
                "service": svc, "correlation_id": cid,
                "received_at": it.get("received_at") or it.get("timestamp") or it.get("created_at"),
                "verdict": _receipt_verdict(it),
                "kind": it.get("kind") or it.get("type"),
                "bundle_ref": f"/v1/receipts/{svc}/{cid}" if cid else None,
            })
    out.sort(key=lambda r: str(r.get("received_at") or ""), reverse=True)
    return {
        "receipts": out[: max(limit, 20)], "count": len(out),
        "services": live, "services_reachable": sum(1 for v in live.values() if v),
        "detail_endpoint": f"{EVIDENCE_RECEIPTS_URL}/v1/receipts/{{service}}/{{correlation_id}}",
    }


_QUERY_LANGS = {"sparql", "cypher", "gremlin"}


class QueryRequest(BaseModel):
    project: str = "default"
    lang: str = "sparql"
    query: str
    params: dict[str, str] | None = None


def _scan_ids(obj: Any, epi: dict[str, str], found: dict[str, str]) -> None:
    """Recursively find any value in a query result that is a known project-entity id, so results can carry the
    fact's epistemic status regardless of the (SPARQL/Cypher/Gremlin) result shape."""
    if isinstance(obj, str):
        if obj in epi:
            found[obj] = epi[obj]
    elif isinstance(obj, dict):
        for v in obj.values():
            _scan_ids(v, epi, found)
    elif isinstance(obj, list):
        for v in obj:
            _scan_ids(v, epi, found)


def _rows_columns(res: dict[str, Any]) -> tuple[list[dict[str, Any]], list[str]]:
    """Best-effort normalise a kernel query result into rows + columns for a grid, defensively across shapes."""
    results = res.get("results")
    if isinstance(results, dict) and isinstance(results.get("bindings"), list):  # SPARQL 1.1 JSON
        cols = list(res.get("head", {}).get("vars", []))
        rows = [{c: (b.get(c, {}) or {}).get("value") for c in cols} for b in results["bindings"]]
        return rows, cols
    rws = res.get("rows")
    if isinstance(rws, list):  # Cypher/Gremlin {columns, rows}
        cols = res.get("columns") or (list(rws[0].keys()) if rws and isinstance(rws[0], dict) else [])
        return rws, list(cols)
    return [], []


@app.post("/api/studio/query")
async def query(req: QueryRequest, _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Proof-carrying query IDE (WS#30). Runs SPARQL / Cypher / Gremlin against the live kernel and returns the
    result WITH its replay proof (query_hash + evaluated_at_seq) AND per-fact epistemic enrichment: every value
    that is a known project entity is tagged with how well it's known. Stardog/Bloom return rows; we return rows
    you can REPLAY and whose facts carry their epistemic status. Bad syntax is a 400, not a silently-empty result."""
    lang = req.lang.lower()
    if lang not in _QUERY_LANGS:
        raise HTTPException(status_code=422, detail=f"lang must be one of {sorted(_QUERY_LANGS)}")
    if not req.query.strip():
        raise HTTPException(status_code=422, detail="query required")
    coll = proj_collection(req.project)
    body: dict[str, Any] = {"query": req.query}
    if lang == "cypher" and req.params:
        body["params"] = req.params
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        (res, err), (subg, _s) = await asyncio.gather(
            _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/{lang}", json=body),
            _req(client, "GET", f"{HELLGRAPH_URL}/api/graph/subgraph?label={coll}&limit=500"),
        )
    if not isinstance(res, dict):
        raise HTTPException(status_code=400, detail=f"query failed: {err or 'no result from kernel'}")
    epi = {n["id"]: n["epistemic_mode"] for n in [_map_node(x) for x in ((subg.get("nodes") if isinstance(subg, dict) else None) or [])]}
    epistemic: dict[str, str] = {}
    _scan_ids(res, epi, epistemic)
    rows, columns = _rows_columns(res)
    proof = {
        "query_hash": res.get("queryHash") or res.get("query_hash"),
        "evaluated_at_seq": res.get("evaluatedAtSeq") or res.get("evaluated_at_seq"),
        "replayable": bool(res.get("queryHash") or res.get("query_hash")),
    }
    return {
        "project": req.project, "lang": lang, "columns": columns, "rows": rows, "row_count": len(rows),
        "epistemic": epistemic, "proof": proof,
        "raw": {k: v for k, v in res.items() if k not in ("queryHash", "evaluatedAtSeq", "ok")},
    }


# ── WS#32: experiment tracking — runs persisted as FIRST-CLASS proof-carrying graph facts (not a side DB) ──────────
def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_json(s: Any) -> Any:
    if isinstance(s, (dict, list)):
        return s
    try:
        return json.loads(s) if s else {}
    except (ValueError, TypeError):
        return {}


async def _fetch_raw_nodes(coll: str, limit: int = 500) -> tuple[list[dict[str, Any]], str | None]:
    """Raw project nodes (properties preserved — unlike _map_node, which projects to a fixed field set)."""
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        res, err = await _req(client, "GET", f"{HELLGRAPH_URL}/api/graph/subgraph?label={coll}&limit={limit}")
    raw = (res.get("nodes") if isinstance(res, dict) else None) or []
    return (raw if isinstance(raw, list) else []), err


class ExperimentRun(BaseModel):
    project: str = "default"
    name: str
    params: dict[str, Any] = {}
    metrics: dict[str, float] = {}
    status: str = "finished"   # running | finished | failed
    source: str | None = None


@app.post("/api/studio/experiments")
async def create_experiment(req: ExperimentRun, authorization: str = Header(default="")) -> dict[str, Any]:
    """Track an experiment RUN. MEET MLflow/W&B: params + metrics + status per run. BEAT: the run is persisted as
    a node in the PROOF-CARRYING graph (labels Run+Experiment), so it isn't a row in a side database — it's a FACT
    carrying epistemic status + provenance, queryable via the query IDE and linkable to the data/models it touched.
    Token-gated (it writes the graph)."""
    _require_write_token(authorization)
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="name required")
    coll = proj_collection(req.project)
    run_id = f"{coll}:run:{uuid.uuid4().hex[:12]}"
    prov = _workbench_prov(coll, "observed", req.source or "studio/experiment")
    prov["extractor"] = "studio/experiment-v0"
    props = {
        "name": name, "run_id": run_id, "status": req.status,
        "params_json": json.dumps(req.params, default=str), "metrics_json": json.dumps(req.metrics, default=str),
        "created_at": _now_iso(), **prov,
    }
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, err = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                            json={"id": run_id, "labels": [coll, "Run", "Experiment"], "properties": props})
    if err:
        raise HTTPException(status_code=502, detail=f"graph write failed: {err}")
    return {"run_id": run_id, "project": req.project, "name": name, "status": req.status,
            "params": req.params, "metrics": req.metrics, "provenance": prov, "written": True}


@app.get("/api/studio/experiments")
async def list_experiments(project: str = "default", limit: int = 200,
                           _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """List the project's experiment runs — read back from the graph, params/metrics parsed, each run carrying its
    epistemic status + provenance. These are graph facts you can also query in the IDE (SPARQL/Cypher over runs)."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, limit)
    runs: list[dict[str, Any]] = []
    for n in raw:
        if "Run" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        runs.append({
            "run_id": p.get("run_id") or n.get("id"), "name": p.get("name"), "status": p.get("status", "unknown"),
            "params": _safe_json(p.get("params_json")), "metrics": _safe_json(p.get("metrics_json")),
            "created_at": p.get("created_at"), "epistemic_mode": p.get("epistemic_mode", "observed"),
            "source": p.get("source"), "extractor": p.get("extractor"),
        })
    runs.sort(key=lambda r: str(r.get("created_at") or ""), reverse=True)
    return {"project": project, "projectCollection": coll, "runs": runs, "count": len(runs), "degraded": err}


# ── WS#35: sovereign persistent IDs + citation. A citable, resolvable identifier for a knowledge artifact —
# DataCite-compatible so it bridges to the scholarly ecosystem — that resolves to a PROOF-CARRYING record
# (provenance + epistemic + content hash), not a bare landing page. The identifier is itself a graph fact. ──
SP_PID_PREFIX = os.getenv("STUDIO_PID_PREFIX", "sp")
STUDIO_DOI_PREFIX = os.getenv("STUDIO_DOI_PREFIX", "10.82044")          # DataCite-style prefix (placeholder until registered)
STUDIO_RESOLVE_BASE = os.getenv("STUDIO_RESOLVE_BASE", "https://studio.socioprophet.ai/resolve")


class CiteRequest(BaseModel):
    project: str = "default"
    kind: str = "graph"        # graph | run | dataset | fact | document
    ref: str = ""              # target id (run_id / node id / dataset id); "" = the whole-project graph snapshot
    title: str | None = None
    creators: list[str] = []


def _mint_pid(coll: str, kind: str, ref: str) -> tuple[str, str]:
    """Content-addressed sovereign PID: stable per (project, kind, target) so citing the same thing is idempotent."""
    h = hashlib.sha256(f"{coll}:{kind}:{ref}".encode()).hexdigest()[:16]
    return f"{SP_PID_PREFIX}:{coll}/{kind}/{h[:12]}", h


@app.post("/api/studio/cite")
async def cite(req: CiteRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Mint a persistent, citable identifier for a knowledge artifact. MEET Zenodo/DataCite: a DOI + a formatted
    citation + BibTeX + DataCite metadata. BEAT: the PID resolves to a PROOF-CARRYING record (the identifier is
    persisted as a Citation graph fact carrying epistemic status + content hash + provenance), and that
    provenance rides inside the DataCite metadata (nanopublication-style) — not a bare landing page."""
    _require_write_token(authorization)
    coll = proj_collection(req.project)
    pid, h = _mint_pid(coll, req.kind, req.ref)
    doi = f"{STUDIO_DOI_PREFIX}/{coll}.{req.kind}.{h[:8]}"
    title = (req.title or f"{req.kind} · {req.ref or coll}").strip()
    creators = req.creators or ["SocioProphet Knowledge Commons"]
    year = datetime.now(timezone.utc).year
    created = _now_iso()
    resolve_url = f"{STUDIO_RESOLVE_BASE}?pid={pid}"
    prov = _workbench_prov(coll, "attested", "studio/cite")   # a minted, hash-sealed identifier is 'attested'
    prov["extractor"] = "studio/cite-v0"
    node_id = f"{coll}:cite:{h[:12]}"
    props = {"pid": pid, "doi": doi, "kind": req.kind, "target": req.ref or coll, "title": title,
             "creators_json": json.dumps(creators), "resolve": resolve_url, "content_hash": h,
             "created_at": created, **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, err = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                            json={"id": node_id, "labels": [coll, "Citation", "Identifier"], "properties": props})
    if err:
        raise HTTPException(status_code=502, detail=f"graph write failed: {err}")
    citation = f"{', '.join(creators)} ({year}). {title}. SocioProphet Knowledge Commons. {pid} (DOI: {doi})."
    bibtex = ("@misc{" + h[:8] + ",\n  author = {" + " and ".join(creators) + "},\n  title = {" + title
              + "},\n  year = {" + str(year) + "},\n  howpublished = {SocioProphet Knowledge Commons},\n  note = {"
              + pid + "},\n  doi = {" + doi + "}\n}")
    datacite = {
        "id": doi, "type": "dois",
        "attributes": {
            "doi": doi, "titles": [{"title": title}], "creators": [{"name": c} for c in creators],
            "publisher": "SocioProphet Knowledge Commons", "publicationYear": year,
            "types": {"resourceTypeGeneral": "Dataset" if req.kind in ("dataset", "graph") else "Other"},
            "url": resolve_url,
            # the BEAT: provenance + epistemic status ride inside the standard metadata (FAIR+)
            "descriptions": [{"descriptionType": "Other",
                              "description": f"Proof-carrying record. epistemic_mode={prov['epistemic_mode']}; content_hash={h}; provenance={prov['extractor']}."}],
        },
    }
    return {"pid": pid, "doi": doi, "resolve": resolve_url, "content_hash": h, "created_at": created,
            "citation": citation, "bibtex": bibtex, "datacite": datacite, "proof_carrying": True, "node_id": node_id}


@app.get("/api/studio/resolve")
async def resolve(pid: str = "", _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Resolve a sovereign PID to its PROOF-CARRYING record — not a landing page. Reads the Citation fact back
    from the graph and returns the target + provenance + epistemic status + content hash, so the identifier
    resolves to something you can VERIFY (the beat over a DOI landing page)."""
    if not pid:
        raise HTTPException(status_code=422, detail="pid required")
    m = re.match(r"^[^:]+:([^/]+)/", pid)
    if not m:
        raise HTTPException(status_code=422, detail="malformed pid")
    coll = m.group(1)
    raw, err = await _fetch_raw_nodes(coll, 500)
    for n in raw:
        p = n.get("properties") or {}
        if "Citation" in (n.get("labels") or []) and p.get("pid") == pid:
            return {"pid": pid, "found": True, "doi": p.get("doi"), "kind": p.get("kind"), "target": p.get("target"),
                    "title": p.get("title"), "creators": _safe_json(p.get("creators_json")),
                    "created_at": p.get("created_at"), "content_hash": p.get("content_hash"),
                    "provenance": {"epistemic_mode": p.get("epistemic_mode"), "extractor": p.get("extractor"), "source": p.get("source")},
                    "proof_carrying": True}
    return {"pid": pid, "found": False, "degraded": err}


# ── WS#36: immutable preservation + versioning. A tamper-evident, content-addressed snapshot of a knowledge
# artifact, versioned in a chain (each links to its predecessor). Idempotent: re-preserving unchanged state
# returns the existing version. The snapshot carries the provenance chain — an archived fact stays proof-carrying. ──
class PreserveRequest(BaseModel):
    project: str = "default"
    target: str = ""        # "" = the whole-project graph; else a label/id to scope the snapshot
    note: str | None = None


def _state_hash(raw: list[dict[str, Any]], target: str) -> str:
    """A stable content hash over the CURRENT state of the target — the sorted (id, content_hash-of-props) of the
    nodes in scope. Deterministic, so identical state → identical hash (tamper-evidence + idempotent versioning)."""
    def in_scope(n: dict[str, Any]) -> bool:
        if not target:
            return "Snapshot" not in (n.get("labels") or [])   # don't hash prior snapshots into the state
        return target in (n.get("labels") or []) or n.get("id") == target
    parts = []
    for n in sorted((x for x in raw if in_scope(x)), key=lambda n: str(n.get("id", ""))):
        props = json.dumps(n.get("properties") or {}, sort_keys=True, default=str)
        parts.append(f"{n.get('id','')}={hashlib.sha256(props.encode()).hexdigest()[:16]}")
    return hashlib.sha256("|".join(parts).encode()).hexdigest()


@app.post("/api/studio/preserve")
async def preserve(req: PreserveRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Seal an immutable, versioned snapshot of a knowledge artifact. MEET Zenodo/OSF versioning. BEAT: the
    snapshot is content-addressed + tamper-evident (re-hash to verify), chained to its predecessor, and carries
    the provenance chain (epistemic=attested) so an archived fact stays proof-carrying. Idempotent — preserving
    unchanged state returns the existing version, never a duplicate."""
    _require_write_token(authorization)
    coll = proj_collection(req.project)
    target = req.target or coll
    raw, err = await _fetch_raw_nodes(coll, 1000)
    if err:
        raise HTTPException(status_code=502, detail=f"graph read failed: {err}")
    content_hash = _state_hash(raw, req.target)
    snaps = [n for n in raw if "Snapshot" in (n.get("labels") or []) and (n.get("properties") or {}).get("target") == target]
    snaps.sort(key=lambda n: (n.get("properties") or {}).get("version", 0))
    # idempotent: unchanged state → return the existing head version
    if snaps and (snaps[-1].get("properties") or {}).get("content_hash") == content_hash:
        p = snaps[-1]["properties"]
        return {"snapshot_id": snaps[-1].get("id"), "version": p.get("version"), "content_hash": content_hash,
                "target": target, "sealed_at": p.get("sealed_at"), "unchanged": True, "proof_carrying": True}
    version = len(snaps) + 1
    parent = snaps[-1].get("id") if snaps else None
    sealed_at = _now_iso()
    snap_id = f"{coll}:snap:{content_hash[:12]}"
    prov = _workbench_prov(coll, "attested", "studio/preserve")
    prov["extractor"] = "studio/preserve-v0"
    props = {"target": target, "version": version, "content_hash": content_hash, "sealed_at": sealed_at,
             "parent": parent or "", "note": req.note or "", **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": snap_id, "labels": [coll, "Snapshot", "Preservation"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        if parent:  # chain to predecessor (version lineage)
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                       json={"label": "supersedes", "from": snap_id, "to": parent, "properties": prov})
    return {"snapshot_id": snap_id, "version": version, "content_hash": content_hash, "target": target,
            "sealed_at": sealed_at, "parent": parent, "unchanged": False, "proof_carrying": True}


@app.get("/api/studio/versions")
async def versions(project: str = "default", target: str = "",
                   _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The version history of a preserved artifact — the immutable chain, newest first, each sealed with its
    content hash + timestamp. Verify integrity by re-hashing the state against a version's content_hash."""
    coll = proj_collection(project)
    tgt = target or coll
    raw, err = await _fetch_raw_nodes(coll, 1000)
    snaps = []
    for n in raw:
        if "Snapshot" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        if p.get("target") != tgt:
            continue
        snaps.append({"snapshot_id": n.get("id"), "version": p.get("version"), "content_hash": p.get("content_hash"),
                      "sealed_at": p.get("sealed_at"), "parent": p.get("parent") or None, "note": p.get("note") or None,
                      "epistemic_mode": p.get("epistemic_mode", "attested")})
    snaps.sort(key=lambda s: s.get("version") or 0, reverse=True)
    return {"project": project, "target": tgt, "versions": snaps, "count": len(snaps), "degraded": err}


# ── WS#37: FAIR metadata + interoperability. Assembles a FAIR record (Findable/Accessible/Interoperable/
# Reusable) from the artifact's citation + preservation, emits schema.org/Dataset JSON-LD (Google Dataset
# Search) alongside DataCite + the PROV-O Turtle export, and scores FAIR. The BEAT is FAIR+: epistemic status
# and a verifiable provenance chain ride inside the metadata (nanopublication-style), which FAIR alone omits. ──
@app.get("/api/studio/fair")
async def fair(project: str = "default", target: str = "",
               _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """FAIR metadata + self-assessment for a knowledge artifact. MEET Zenodo/OpenAIRE: schema.org/Dataset JSON-LD +
    DataCite + PROV-O Turtle + an F/A/I/R score. BEAT (FAIR+): epistemic status + a verifiable provenance chain
    ride inside the record — Interoperability is real (RDF/PROV-O export), not just 'a metadata blob exists'."""
    coll = proj_collection(project)
    tgt = target or coll
    raw, err = await _fetch_raw_nodes(coll, 1000)
    cite = next((n.get("properties") or {} for n in raw
                 if "Citation" in (n.get("labels") or []) and (n.get("properties") or {}).get("target") == tgt), {})
    snaps = [n.get("properties") or {} for n in raw
             if "Snapshot" in (n.get("labels") or []) and (n.get("properties") or {}).get("target") == tgt]
    snaps.sort(key=lambda p: p.get("version", 0))
    pid, doi = cite.get("pid"), cite.get("doi")
    title = cite.get("title") or f"{project} — knowledge graph"
    resolve = cite.get("resolve")
    version = snaps[-1].get("version") if snaps else None
    content_hash = (snaps[-1].get("content_hash") if snaps else None) or cite.get("content_hash")
    turtle_export = f"/api/studio/graph.ttl?project={project}"
    schema_org = {
        "@context": "https://schema.org/", "@type": "Dataset", "name": title,
        "identifier": doi or pid, "url": resolve,
        "creator": {"@type": "Organization", "name": "SocioProphet Knowledge Commons"},
        "distribution": {"@type": "DataDownload", "encodingFormat": "text/turtle", "contentUrl": turtle_export},
        "version": version, "sha256": content_hash,
        # FAIR+: provenance + epistemic status ride in the record
        "provenance": {"epistemic_status": cite.get("epistemic_mode", "observed"), "hashSealed": bool(content_hash)},
    }
    findable, accessible = bool(pid or doi), bool(resolve)
    interoperable = True                      # RDF/Turtle + PROV-O + schema.org always available
    reusable = bool(cite)                     # provenance present (+ license, implicit)
    score = round(sum([findable, accessible, interoperable, reusable]) / 4, 2)
    return {
        "project": project, "target": tgt, "title": title, "pid": pid, "doi": doi,
        "version": version, "content_hash": content_hash,
        "schema_org": schema_org, "datacite_doi": doi, "turtle_export": turtle_export,
        "fair": {"findable": findable, "accessible": accessible, "interoperable": interoperable, "reusable": reusable, "score": score},
        "fair_plus": {"epistemic": True, "provenance_chain": bool(snaps), "hash_sealed": bool(content_hash)},
        "hint": None if findable else "mint a persistent identifier (Cite) to make this Findable",
        "degraded": err,
    }


# ── WS#38: scholarly + agent ecosystem hooks. MEET Zenodo/OpenAIRE: ORCID contributor links, DOI resolution,
# an OpenAIRE/DataCite-harvestable metadata pointer. BEAT (agent-native): a machine-readable capability manifest
# so an AGENT — not just a human — can discover the proof-carrying record and consume it, each access verb flagged
# verifiable. Scholarly repos expose discovery to people; we expose it to autonomous agents. ────────────────────
@app.get("/api/studio/ecosystem")
async def ecosystem(project: str = "default", target: str = "",
                    _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Scholarly + agent ecosystem hooks for a knowledge artifact: ORCID contributors, DOI resolution, an
    OpenAIRE-harvestable pointer, and an agent-native capability manifest of verifiable access verbs."""
    coll = proj_collection(project)
    tgt = target or coll
    raw, err = await _fetch_raw_nodes(coll, 1000)
    cite = next((n.get("properties") or {} for n in raw
                 if "Citation" in (n.get("labels") or []) and (n.get("properties") or {}).get("target") == tgt), {})
    contributors: list[dict[str, Any]] = []
    seen: set[str] = set()
    for n in raw:                                   # harvest ORCID from any graph node that carries one
        p = n.get("properties") or {}
        orcid = p.get("orcid")
        if orcid and orcid not in seen:
            seen.add(orcid)
            contributors.append({"name": p.get("name") or p.get("label") or n.get("id"),
                                 "orcid": orcid, "orcid_url": f"https://orcid.org/{orcid}"})
    for c in (cite.get("contributors") or []):      # plus any the citation declared
        oc = c.get("orcid") if isinstance(c, dict) else None
        if oc and oc not in seen:
            seen.add(oc)
            contributors.append({"name": c.get("name"), "orcid": oc, "orcid_url": f"https://orcid.org/{oc}"})
    doi, pid = cite.get("doi"), cite.get("pid")
    scholarly = {
        "doi": doi, "doi_url": f"https://doi.org/{doi}" if doi else None,
        "orcid_contributors": contributors,
        "openaire": {"harvestable": bool(doi), "datacite_doi": doi,
                     "metadata": f"/api/studio/fair?project={project}&target={tgt}"},
    }

    def _verb(name: str, endpoint: str | None) -> dict[str, Any]:
        return {"name": name, "endpoint": endpoint, "verifiable": True}

    agent_manifest = {
        "@type": "AgentCapabilityManifest", "commons": "SocioProphet Knowledge Commons",
        "project": project, "identifier": pid or doi,
        "proof_carrying": True, "epistemic_status": True, "sovereign": True,
        "access": [
            _verb("resolve", f"/api/studio/resolve?pid={pid}" if pid else None),
            _verb("query", "/api/studio/query"),
            _verb("provenance", "/api/studio/provenance"),
            _verb("receipts", "/api/studio/receipts"),
            _verb("rdf", f"/api/studio/graph.ttl?project={project}"),
            _verb("fair", f"/api/studio/fair?project={project}&target={tgt}"),
        ],
        "consume_note": "every result carries a queryHash + epistemic status; verify via the receipts/provenance verbs",
    }
    return {"project": project, "target": tgt, "scholarly": scholarly,
            "agent_manifest": agent_manifest, "degraded": err}


# ── WS#43: the research/PID graph over HellGraph. MEET OpenAIRE Research Graph / DataCite Commons: a typed graph
# connecting Results (cited artifacts) ⇄ Persons (ORCIDs) ⇄ Organizations ⇄ Identifiers (PIDs/DOIs). We build it
# NATIVELY over HellGraph from the citations (WS#35) + contributors (WS#38) already in the project collection —
# no separate metadata store. BEAT: every node and edge is PROOF-CARRYING (epistemic status + provenance), which
# a bare bibliographic graph is not; and it's live-queryable in the same kernel as the knowledge itself. ────────
@app.get("/api/studio/pidgraph")
async def pidgraph(project: str = "default",
                   _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The project's research/PID graph — Results ⇄ Persons(ORCID) ⇄ Organizations ⇄ Identifiers(PID/DOI),
    assembled from Citation + contributor nodes. The OpenAIRE/DataCite-Commons pattern, native + proof-carrying."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 1000)
    nodes: dict[str, dict[str, Any]] = {}
    edges: list[dict[str, Any]] = []

    def _add(nid: str, ntype: str, label: str, **extra: Any) -> None:
        if nid and nid not in nodes:
            nodes[nid] = {"id": nid, "type": ntype, "label": label, "proof_carrying": True, **extra}

    def _edge(frm: str, to: str, label: str) -> None:
        if frm and to:
            edges.append({"from": frm, "to": to, "label": label})

    for n in raw:
        if "Citation" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        target = p.get("target") or n.get("id")
        pid, doi = p.get("pid"), p.get("doi")
        epi = p.get("epistemic_mode", "observed")
        # the Result (the cited artifact) + its Identifier(s)
        _add(target, "Result", p.get("title") or target, epistemic_mode=epi)
        if pid:
            _add(pid, "Identifier", pid, scheme="sovereign-pid", epistemic_mode=epi)
            _edge(pid, target, "identifies")
        if doi:
            _add(doi, "Identifier", doi, scheme="datacite-doi", epistemic_mode=epi)
            _edge(doi, target, "identifies")
        # Persons (ORCID contributors) authored the Result; Organizations they name
        for c in (p.get("contributors") or []):
            if not isinstance(c, dict):
                continue
            orcid = c.get("orcid")
            person = f"orcid:{orcid}" if orcid else f"person:{_norm(c.get('name') or '')}"
            _add(person, "Person", c.get("name") or orcid or person, orcid=orcid, epistemic_mode="attested")
            _edge(person, target, "authored")
            org = c.get("affiliation") or c.get("org")
            if org:
                oid = f"org:{_norm(org)}"
                _add(oid, "Organization", org, epistemic_mode="attested")
                _edge(person, oid, "affiliated")
    # Persons declared as standalone graph nodes (an orcid property anywhere) also join the graph
    for n in raw:
        p = n.get("properties") or {}
        orcid = p.get("orcid")
        if orcid:
            _add(f"orcid:{orcid}", "Person", p.get("name") or orcid, orcid=orcid,
                 epistemic_mode=p.get("epistemic_mode", "attested"))

    counts: dict[str, int] = {}
    for nd in nodes.values():
        counts[nd["type"]] = counts.get(nd["type"], 0) + 1
    return {
        "project": project, "collection": coll,
        "nodes": list(nodes.values()), "edges": edges,
        "stats": {"nodes": len(nodes), "edges": len(edges),
                  "results": counts.get("Result", 0), "persons": counts.get("Person", 0),
                  "organizations": counts.get("Organization", 0), "identifiers": counts.get("Identifier", 0)},
        "pattern": "OpenAIRE Research Graph / DataCite Commons — built natively over HellGraph",
        "beat": "every node & edge is proof-carrying (epistemic status + provenance) and live-queryable — not a bare bibliographic graph",
        "degraded": err,
    }


# ── WS#39: commons at scale + community curation. MEET Zenodo/Wikidata: a commons overview (scale + community
# stats) and community endorsement of records. BEAT (epistemic curation): endorsements are governed, proof-
# carrying facts (identified endorser, revocable), and the curation score is EPISTEMIC-WEIGHTED — trust follows
# the epistemic status of the underlying facts (attested > verified > observed …), not raw popularity. ──────────
# epistemic ladder → curation weight (higher = more trustworthy grounding). Mirrors studioApi EPISTEMIC_ORDER.
EPISTEMIC_WEIGHT = {"attested": 1.0, "verified": 0.85, "observed": 0.6,
                    "derived": 0.45, "hypothesis": 0.25, "simulated": 0.1}


@app.get("/api/studio/commons")
async def commons(project: str = "default",
                  _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Commons-at-scale overview for a project: node/edge counts, the epistemic-status distribution, and the
    community signals (citations, preserved versions, endorsements, contributors). The scale + health story —
    every count is grounded in the graph, and the epistemic distribution is the quality signal repos lack."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    epistemic: dict[str, int] = {}
    citations = versions = endorsements = 0
    contributors: set[str] = set()
    facts = 0
    for n in raw:
        labels = n.get("labels") or []
        p = n.get("properties") or {}
        if "Citation" in labels:
            citations += 1
        elif "Snapshot" in labels:
            versions += 1
        elif "Endorsement" in labels:
            endorsements += 1
        else:
            facts += 1
            mode = p.get("epistemic_mode")
            if mode:
                epistemic[mode] = epistemic.get(mode, 0) + 1
        if p.get("orcid"):
            contributors.add(p["orcid"])
    # a single epistemic-weighted quality index over the project's facts (0..1)
    graded = sum(EPISTEMIC_WEIGHT.get(m, 0.3) * c for m, c in epistemic.items())
    total_graded = sum(epistemic.values())
    quality_index = round(graded / total_graded, 3) if total_graded else None
    return {
        "project": project, "collection": coll,
        "scale": {"facts": facts, "citations": citations, "preserved_versions": versions,
                  "endorsements": endorsements, "contributors": len(contributors)},
        "epistemic_distribution": epistemic,
        "epistemic_quality_index": quality_index,   # the beat: quality, not just volume
        "degraded": err,
    }


class EndorseRequest(BaseModel):
    project: str = "default"
    target: str            # node id (or label) being endorsed
    endorser: str          # identified endorser (orcid / sovereign id / handle) — no anonymous curation
    note: str | None = None
    revoke: bool = False    # governed: an endorsement can be withdrawn


@app.post("/api/studio/endorse")
async def endorse(req: EndorseRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Community curation: an identified endorser endorses (or revokes) a record. BEAT: the endorsement is a
    governed, proof-carrying graph fact — identified endorser, timestamped, revocable — not an anonymous vote.
    Idempotent per (target, endorser): re-endorsing updates; revoke marks it withdrawn."""
    _require_write_token(authorization)
    if not req.target.strip() or not req.endorser.strip():
        raise HTTPException(status_code=422, detail="target and endorser required")
    coll = proj_collection(req.project)
    endorser_key = hashlib.sha256(req.endorser.encode()).hexdigest()[:12]
    end_id = f"{coll}:endorse:{hashlib.sha256(req.target.encode()).hexdigest()[:8]}:{endorser_key}"
    prov = _workbench_prov(coll, "attested", "studio/endorse")
    prov["extractor"] = "lattice-studio/endorse-v0"
    props = {"target": req.target, "endorser": req.endorser, "note": req.note or "",
             "revoked": req.revoke, "at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": end_id, "labels": [coll, "Endorsement", "Curation"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        if not req.revoke:      # link the endorsement to what it endorses (revoked ones stay as tombstones)
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                       json={"label": "endorses", "from": end_id, "to": req.target, "properties": prov})
    return {"endorsement_id": end_id, "target": req.target, "endorser": req.endorser,
            "revoked": req.revoke, "proof_carrying": True}


@app.get("/api/studio/curation")
async def curation(project: str = "default", target: str = "",
                   _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The curation signals for a target: the identified, non-revoked endorsements and an epistemic-weighted
    curation score. BEAT: the score weights each endorsement by the epistemic status of the endorsed fact —
    an endorsement of an attested fact counts more than one of a hypothesis. Governance is read-enforced:
    revoked endorsements never contribute."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    by_id = {n.get("id"): (n.get("properties") or {}) for n in raw}
    fact_mode = {nid: p.get("epistemic_mode") for nid, p in by_id.items()}
    endorsements = []
    seen: set[str] = set()
    for n in raw:
        if "Endorsement" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        if target and p.get("target") != target:
            continue
        if p.get("revoked"):        # governance: withdrawn curation never counts
            continue
        key = f"{p.get('target')}|{p.get('endorser')}"
        if key in seen:
            continue
        seen.add(key)
        endorsements.append({"target": p.get("target"), "endorser": p.get("endorser"),
                             "note": p.get("note") or None, "at": p.get("at")})
    # epistemic-weighted curation score: each endorsement scaled by the grounding of the fact it endorses
    score = 0.0
    for e in endorsements:
        mode = fact_mode.get(e["target"]) or "observed"
        score += EPISTEMIC_WEIGHT.get(mode, 0.3)
    return {"project": project, "target": target or None, "endorsements": endorsements,
            "count": len(endorsements), "curation_score": round(score, 3),
            "epistemic_weighted": True, "degraded": err}


# ── WS#33: data connectors framework (governed ingest). MEET DS studios: connectors that pull structured data
# into the workspace. BEAT: ingest is GOVERNED — fail-closed write gate, per-row provenance, epistemic status,
# and a source connector id — so every ingested row lands as a proof-carrying fact, not an untracked blob. Inline
# CSV/JSON is fully live here; fetch-based connectors (http/s3/postgres) are declared + gated, not yet wired. ──
INGEST_ROW_CAP = 5000

CONNECTORS = [
    {"type": "csv", "status": "live", "governed": True, "note": "inline CSV → one proof-carrying node per row"},
    {"type": "json", "status": "live", "governed": True, "note": "inline JSON array of objects → node per element"},
    {"type": "http", "status": "declared", "governed": True, "note": "URL fetch — membrane-gated egress, not yet wired"},
    {"type": "s3", "status": "declared", "governed": True, "note": "object store — sovereign creds required, not yet wired"},
    {"type": "postgres", "status": "declared", "governed": True, "note": "SQL source — governed pull, not yet wired"},
]


# The fetch/SaaS-connector backbone: oomol-lab/open-connector (Apache-2.0, agent-native — SDK/CLI/MCP/HTTP/
# OpenAPI, 1000+ providers, self-hostable). We consume it sovereignly behind our governance gate — connections
# are registered as proof-carrying facts here; the byte-level auth + fetch executes in the open-connector service.
OPEN_CONNECTOR = {"project": "oomol-lab/open-connector", "license": "Apache-2.0",
                  "role": "fetch/SaaS backbone", "interfaces": ["sdk", "cli", "mcp", "http", "openapi"],
                  "url": "https://github.com/oomol-lab/open-connector"}


@app.get("/api/studio/connectors")
async def connectors(_auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The connector registry: supported source types and each one's governance posture. Inline csv/json are live;
    fetch/SaaS connectors run on the open-connector backbone, registered here as governed, proof-carrying facts."""
    return {"connectors": CONNECTORS, "row_cap": INGEST_ROW_CAP, "backbone": OPEN_CONNECTOR,
            "governance": "fail-closed write token + per-row provenance + epistemic status on every ingested fact"}


# ── WS#41: governed connection registry (over the open-connector backbone). MEET Databricks/Foundry connection
# management: register a source connection to a SaaS/DB/object-store provider. BEAT: a connection is a governed,
# proof-carrying graph fact — provider + owner + status, revocable — and any ingest through it carries the
# connection's provenance. Byte-level auth + fetch executes in the open-connector service (deploy-gated). ────────
class ConnectRequest(BaseModel):
    project: str = "default"
    provider: str                          # e.g. github, gmail, bigquery, s3, postgres — the open-connector catalog
    name: str | None = None                # a human label for this connection
    note: str | None = None
    owner: str | None = None               # who owns/authorised the connection (provenance)


@app.post("/api/studio/connect")
async def connect(req: ConnectRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Register a governed connection to a provider (on the open-connector backbone). Fail-closed; idempotent per
    (project, provider, name). Status is 'declared' until the open-connector service completes the OAuth/token
    handshake — the registry entry is a proof-carrying fact regardless, so ingest can carry its provenance."""
    _require_write_token(authorization)
    provider = req.provider.strip().lower()
    if not provider:
        raise HTTPException(status_code=422, detail="provider required")
    coll = proj_collection(req.project)
    label = (req.name or provider).strip()
    cid = f"{coll}:connection:{_norm(provider)}:{_norm(label).replace(' ', '_')}"
    prov = _workbench_prov(coll, "attested", req.owner or "studio/connect")
    prov["extractor"] = "lattice-studio/connection-v0"
    props = {"provider": provider, "name": label, "status": "declared", "backbone": "open-connector",
             "owner": req.owner or "", "note": req.note or "", "created_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": cid, "labels": [coll, "Connection"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
    return {"connection_id": cid, "provider": provider, "name": label, "status": "declared",
            "backbone": "oomol-lab/open-connector", "proof_carrying": True,
            "note": "governed connection registered; OAuth/token handshake + fetch execute in the open-connector service"}


@app.get("/api/studio/connections")
async def connections(project: str = "default",
                      _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The project's governed connections — provider, owner, status — each a proof-carrying, revocable graph fact."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = []
    for n in raw:
        if "Connection" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        out.append({"connection_id": n.get("id"), "provider": p.get("provider"), "name": p.get("name"),
                    "status": p.get("status", "declared"), "owner": p.get("owner") or None,
                    "backbone": p.get("backbone", "open-connector"), "created_at": p.get("created_at")})
    out.sort(key=lambda x: x.get("created_at") or "", reverse=True)
    return {"project": project, "connections": out, "count": len(out),
            "backbone": OPEN_CONNECTOR, "degraded": err}


class IngestRequest(BaseModel):
    project: str = "default"
    connector: str = "csv"          # csv | json (live); others rejected until wired
    data: str = ""                  # inline CSV text, or a JSON array-of-objects string
    key: str | None = None          # column/field to use as the node key (dedup); else row index
    label: str = "Record"           # graph label for ingested rows
    epistemic_mode: str = "observed"
    source: str | None = None


def _parse_rows(connector: str, data: str) -> list[dict[str, Any]]:
    """Parse inline CSV or JSON-array-of-objects into a list of flat row dicts. Deterministic + real — no
    external I/O. Raises HTTPException(422) on an unusable payload or an unsupported (declared-only) connector."""
    if connector == "csv":
        rows = list(csv.DictReader(io.StringIO(data)))
        if not rows:
            raise HTTPException(status_code=422, detail="csv payload has no data rows")
        return [dict(r) for r in rows]
    if connector == "json":
        try:
            parsed = json.loads(data)
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=422, detail=f"invalid json: {e}") from e
        if not isinstance(parsed, list) or not all(isinstance(x, dict) for x in parsed):
            raise HTTPException(status_code=422, detail="json connector expects an array of objects")
        if not parsed:
            raise HTTPException(status_code=422, detail="json payload is empty")
        return parsed
    raise HTTPException(status_code=422, detail=f"connector '{connector}' is declared but not yet wired for ingest")


@app.post("/api/studio/ingest")
async def ingest(req: IngestRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Governed ingest of inline structured data (csv/json) into the project graph. Each row becomes a proof-
    carrying node: per-row provenance (source connector + epistemic status), fail-closed write gate, row cap,
    dedup by the key column. The BEAT over a plain DS connector: every ingested fact is governed + attributable."""
    _require_write_token(authorization)
    rows = _parse_rows(req.connector, req.data)
    if len(rows) > INGEST_ROW_CAP:
        raise HTTPException(status_code=413, detail=f"ingest exceeds row cap ({len(rows)} > {INGEST_ROW_CAP})")
    coll = proj_collection(req.project)
    src = req.source or f"connector:{req.connector}:{hashlib.sha256(req.data.encode()).hexdigest()[:12]}"
    prov = _workbench_prov(coll, req.epistemic_mode, src)
    prov["extractor"] = f"lattice-studio/connector-{req.connector}-v0"
    prov["connector"] = req.connector

    written = 0
    errors: list[str] = []
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        calls = []
        for i, row in enumerate(rows):
            key_val = str(row.get(req.key)) if req.key and row.get(req.key) is not None else f"row{i}"
            nid = f"{coll}:ingest:{_norm(key_val).replace(' ', '_')}"
            props = {**{str(k): v for k, v in row.items()}, **prov, "ingest_key": key_val}
            calls.append(_req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                              json={"id": nid, "labels": [coll, req.label, "Ingested"], "properties": props}))
        for _, err in await asyncio.gather(*calls):
            if err:
                errors.append(err)
            else:
                written += 1
    return {"project": req.project, "connector": req.connector, "source": src,
            "rows": len(rows), "written": written, "label": req.label,
            "provenance": {"epistemic_mode": req.epistemic_mode, "connector": req.connector, "source": src},
            "errors": errors[:5] or None}


# ── WS#34: explorer UX parity+ — saved perspectives. MEET the studios: named, reusable explorer views (a saved
# label/epistemic filter + layout). BEAT: a perspective is a proof-carrying graph fact — it's shared to the agent
# team the moment it's saved (lives in the project collection), governed (fail-closed write), and can filter the
# explorer BY EPISTEMIC STATUS — the moat, made a first-class lens over the graph. ──────────────────────────────
class PerspectiveRequest(BaseModel):
    project: str = "default"
    name: str
    label: str | None = None            # graph label to focus (e.g. "Person")
    epistemic: list[str] | None = None  # epistemic modes to include (the beat: filter by grounding)
    limit: int = 300
    layout: str = "force"               # force | radial | hierarchy (explorer layout)
    note: str | None = None


@app.post("/api/studio/perspective")
async def save_perspective(req: PerspectiveRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Save a named explorer perspective (label + epistemic filter + layout) as a proof-carrying graph fact —
    shared to the agent team on save, governed, and re-openable. Idempotent per (project, name)."""
    _require_write_token(authorization)
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="name required")
    coll = proj_collection(req.project)
    pid = f"{coll}:perspective:{_norm(name).replace(' ', '_')}"
    prov = _workbench_prov(coll, "attested", "studio/perspective")
    prov["extractor"] = "lattice-studio/perspective-v0"
    props = {"name": name, "label": req.label or "", "epistemic": json.dumps(req.epistemic or []),
             "limit": req.limit, "layout": req.layout, "note": req.note or "", "saved_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": pid, "labels": [coll, "Perspective", "Curation"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
    return {"perspective_id": pid, "name": name, "shared_to_team": True, "proof_carrying": True}


@app.get("/api/studio/perspectives")
async def perspectives(project: str = "default",
                       _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The saved explorer perspectives for a project — each a proof-carrying, team-shared view with its label +
    epistemic filter + layout, ready to re-open in the graph explorer."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 1000)
    out = []
    for n in raw:
        if "Perspective" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        try:
            epi = json.loads(p.get("epistemic") or "[]")
        except (ValueError, TypeError):
            epi = []
        out.append({"perspective_id": n.get("id"), "name": p.get("name"), "label": p.get("label") or None,
                    "epistemic": epi, "limit": p.get("limit", 300), "layout": p.get("layout", "force"),
                    "note": p.get("note") or None, "saved_at": p.get("saved_at")})
    out.sort(key=lambda x: x.get("saved_at") or "", reverse=True)
    return {"project": project, "perspectives": out, "count": len(out), "degraded": err}


async def _fetch_graph(coll: str, limit: int = 2000) -> tuple[list[dict[str, Any]], list[dict[str, Any]], str | None]:
    """Project subgraph as (nodes, edges, err). Edges carry {from, to, label} — used for lineage walks."""
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        res, err = await _req(client, "GET", f"{HELLGRAPH_URL}/api/graph/subgraph?label={coll}&limit={limit}")
    nodes = (res.get("nodes") if isinstance(res, dict) else None) or []
    edges = (res.get("edgeList") if isinstance(res, dict) else None) or []
    return (nodes if isinstance(nodes, list) else []), (edges if isinstance(edges, list) else []), err


def _node_type(n: dict[str, Any], coll: str) -> str:
    return next((l for l in (n.get("labels") or []) if l != coll), "Node")


# ── WS#45: pipelines / workflows. MEET Databricks Workflows + Foundry Pipeline Builder: define a DAG of steps
# (extract → transform → train → …) and record runs. BEAT: the pipeline, every step, and every run are proof-
# carrying HellGraph facts with NATIVE lineage (feeds / step_of / run_of edges) + epistemic status — the lineage
# is the graph, not a bolt-on. Governed (fail-closed writes). Step EXECUTION runs on the notebook/Ray runtime. ──
class PipelineStep(BaseModel):
    id: str
    kind: str = "transform"                 # extract | transform | train | evaluate | publish
    inputs: list[str] = []
    outputs: list[str] = []
    note: str | None = None


class PipelineRequest(BaseModel):
    project: str = "default"
    name: str
    steps: list[PipelineStep] = []


def _pipe_id(coll: str, name: str) -> str:
    return f"{coll}:pipeline:{_norm(name).replace(' ', '_')}"


@app.post("/api/studio/pipeline")
async def upsert_pipeline(req: PipelineRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Define a pipeline as a proof-carrying DAG: a Pipeline node + a Step node per step, with feeds edges wired
    from each step's inputs to the step that produced them, and step_of edges to the pipeline."""
    _require_write_token(authorization)
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="name required")
    coll = proj_collection(req.project)
    pid = _pipe_id(coll, name)
    prov = _workbench_prov(coll, "attested", "studio/pipeline")
    prov["extractor"] = "lattice-studio/pipeline-v0"
    producer: dict[str, str] = {}                         # output name → step id that produces it
    for s in req.steps:
        for o in s.outputs:
            producer[o] = s.id
    props = {"name": name, "steps": json.dumps([s.model_dump() for s in req.steps]),
             "step_count": len(req.steps), "updated_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": pid, "labels": [coll, "Pipeline"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        for s in req.steps:
            sid = f"{pid}:step:{_norm(s.id).replace(' ', '_')}"
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                       json={"id": sid, "labels": [coll, "Step", s.kind.capitalize()],
                             "properties": {"step": s.id, "kind": s.kind, "note": s.note or "", **prov}})
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                       json={"label": "step_of", "from": sid, "to": pid, "properties": prov})
            for inp in s.inputs:                          # wire lineage: producer step → this step
                if inp in producer and producer[inp] != s.id:
                    psid = f"{pid}:step:{_norm(producer[inp]).replace(' ', '_')}"
                    await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                               json={"label": "feeds", "from": psid, "to": sid, "properties": prov})
    return {"pipeline_id": pid, "name": name, "steps": len(req.steps), "proof_carrying": True}


@app.get("/api/studio/pipelines")
async def pipelines(project: str = "default",
                    _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The project's pipelines, each with its step DAG."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 1000)
    out = []
    for n in raw:
        if "Pipeline" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        try:
            steps = json.loads(p.get("steps") or "[]")
        except (ValueError, TypeError):
            steps = []
        out.append({"pipeline_id": n.get("id"), "name": p.get("name"), "steps": steps,
                    "step_count": p.get("step_count", len(steps)), "updated_at": p.get("updated_at")})
    out.sort(key=lambda x: x.get("updated_at") or "", reverse=True)
    return {"project": project, "pipelines": out, "count": len(out), "degraded": err}


class PipelineRunRequest(BaseModel):
    project: str = "default"
    pipeline: str
    status: str = "finished"                 # running | finished | failed
    note: str | None = None


@app.post("/api/studio/pipeline/run")
async def run_pipeline(req: PipelineRunRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Record a pipeline run as a proof-carrying fact (PipelineRun+Run node, run_of edge to the pipeline). This is
    the governed RUN LEDGER; actual step execution runs on the notebook/Ray runtime and emits its own receipts."""
    _require_write_token(authorization)
    coll = proj_collection(req.project)
    pipe_id = _pipe_id(coll, req.pipeline)
    prov = _workbench_prov(coll, "attested", "studio/pipeline-run")
    prov["extractor"] = "lattice-studio/pipeline-run-v0"
    rid = f"{coll}:pipelinerun:{hashlib.sha256((req.pipeline + _now_iso()).encode()).hexdigest()[:12]}"
    props = {"pipeline": req.pipeline, "status": req.status, "note": req.note or "", "ran_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": rid, "labels": [coll, "PipelineRun", "Run"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "run_of", "from": rid, "to": pipe_id, "properties": prov})
    return {"run_id": rid, "pipeline": req.pipeline, "status": req.status, "proof_carrying": True,
            "note": "run ledger recorded as a proof-carrying fact; step execution runs on the notebook/Ray runtime"}


# ── WS#46: data catalog + end-to-end lineage. MEET Databricks Unity Catalog + Foundry datasets/ontology: a
# governed dataset catalog and column/asset lineage. BEAT: datasets are proof-carrying graph nodes (provenance +
# epistemic status native), and lineage is ONE unified graph walk across data → pipeline → run → model → citation,
# not a separate lineage service stitched on afterward. ─────────────────────────────────────────────────────────
_CATALOG_RESERVED = {"epistemic_mode", "source", "extractor", "project", "kko_type", "connector",
                     "ingest_key", "name", "title", "updated_at"}


@app.get("/api/studio/catalog")
async def catalog(project: str = "default",
                  _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The governed dataset catalog: every ingested/declared dataset with its (best-effort) columns, source
    connector, provenance and epistemic status. Every entry is a proof-carrying graph node, not a catalog row."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = []
    for n in raw:
        labels = set(n.get("labels") or [])
        if not ({"Dataset", "Ingested"} & labels):
            continue
        p = n.get("properties") or {}
        cols = [k for k in p.keys() if k not in _CATALOG_RESERVED][:24]
        out.append({"id": n.get("id"), "name": p.get("name") or p.get("title") or n.get("id"),
                    "labels": [l for l in (n.get("labels") or []) if l != coll],
                    "connector": p.get("connector"), "source": p.get("source"),
                    "epistemic_mode": p.get("epistemic_mode", "observed"), "governed": True, "columns": cols})
    return {"project": project, "datasets": out, "count": len(out),
            "beat": "datasets are proof-carrying graph nodes — provenance + epistemic status are native, not a bolt-on catalog",
            "degraded": err}


_LINEAGE_EDGES = {"feeds", "run_of", "step_of", "produced_by", "produces", "derived_from", "supersedes", "in_pipeline"}


@app.get("/api/studio/lineage")
async def lineage(project: str = "default", target: str = "", depth: int = 4,
                  _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """End-to-end lineage DAG for a node — a bounded walk (both directions) over the flow edges (feeds / run_of /
    step_of / produced_by / produces / derived_from / supersedes / in_pipeline) up to `depth`. One unified,
    proof-carrying lineage across data → pipeline → run → model, not a separate lineage service."""
    if not target:
        raise HTTPException(status_code=422, detail="target node id required")
    coll = proj_collection(project)
    nodes, edges, err = await _fetch_graph(coll, 2000)
    by_id = {n.get("id"): n for n in nodes}
    adj: dict[str, list[tuple[str, dict[str, Any]]]] = {}
    for e in edges:
        if (e.get("label") or "") not in _LINEAGE_EDGES:
            continue
        f, t = e.get("from"), e.get("to")
        if not f or not t:
            continue
        adj.setdefault(f, []).append((t, e))
        adj.setdefault(t, []).append((f, e))          # walk both directions
    seen_n: set[str] = {target}
    seen_e: set[tuple[str, str, str]] = set()
    out_edges: list[dict[str, Any]] = []
    frontier = [(target, 0)]
    while frontier:
        nid, d = frontier.pop()
        if d >= max(1, min(depth, 8)):
            continue
        for nb, e in adj.get(nid, []):
            key = (e.get("from"), e.get("to"), e.get("label") or "")
            if key not in seen_e:
                seen_e.add(key)
                out_edges.append({"from": e.get("from"), "to": e.get("to"), "label": e.get("label")})
            if nb not in seen_n:
                seen_n.add(nb)
                frontier.append((nb, d + 1))
    out_nodes = []
    for nid in seen_n:
        n = by_id.get(nid)
        p = (n.get("properties") if n else {}) or {}
        out_nodes.append({"id": nid, "type": _node_type(n, coll) if n else "External",
                          "label": p.get("name") or p.get("title") or nid,
                          "epistemic_mode": p.get("epistemic_mode")})
    return {"project": project, "target": target, "nodes": out_nodes, "edges": out_edges,
            "stats": {"nodes": len(out_nodes), "edges": len(out_edges)},
            "beat": "one proof-carrying lineage graph spanning data, pipelines, runs and models — every hop verifiable",
            "degraded": err}


# ── WS#47: model registry. MEET MLflow / Databricks Model Registry: register model versions, stage them
# (staging → production), track metrics. BEAT: a model version is a proof-carrying graph node linked to the RUN
# that produced it (produced_by edge → WS#32 experiment), so its provenance + epistemic lineage travel with it,
# and a stage transition is a governed, attributable event — not a mutable row. ─────────────────────────────────
_MODEL_STAGES = ["none", "staging", "production", "archived"]


class ModelRequest(BaseModel):
    project: str = "default"
    name: str
    version: str = "1"
    run: str | None = None                   # the experiment run id (WS#32) that produced this version
    stage: str = "none"
    metrics: dict[str, float] = {}
    note: str | None = None


def _model_id(coll: str, name: str, version: str) -> str:
    return f"{coll}:model:{_norm(name).replace(' ', '_')}:{_norm(version)}"


@app.post("/api/studio/model")
async def register_model(req: ModelRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Register a model version as a proof-carrying node, linked (produced_by) to the run that produced it."""
    _require_write_token(authorization)
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="name required")
    if req.stage not in _MODEL_STAGES:
        raise HTTPException(status_code=422, detail=f"stage must be one of {_MODEL_STAGES}")
    coll = proj_collection(req.project)
    mid = _model_id(coll, name, req.version)
    prov = _workbench_prov(coll, "attested", "studio/model")
    prov["extractor"] = "lattice-studio/model-registry-v0"
    props = {"name": name, "version": req.version, "stage": req.stage, "metrics": json.dumps(req.metrics),
             "run": req.run or "", "note": req.note or "", "updated_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": mid, "labels": [coll, "Model"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        if req.run:                              # lineage to the producing experiment run (WS#32)
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                       json={"label": "produced_by", "from": mid, "to": req.run, "properties": prov})
    return {"model_id": mid, "name": name, "version": req.version, "stage": req.stage, "proof_carrying": True}


@app.get("/api/studio/models")
async def models(project: str = "default",
                 _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The model registry: versions grouped by model name, each with stage, metrics and the run it came from."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    grouped: dict[str, list[dict[str, Any]]] = {}
    for n in raw:
        if "Model" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        try:
            metrics = json.loads(p.get("metrics") or "{}")
        except (ValueError, TypeError):
            metrics = {}
        grouped.setdefault(p.get("name") or "?", []).append(
            {"model_id": n.get("id"), "version": p.get("version", "1"), "stage": p.get("stage", "none"),
             "metrics": metrics, "run": p.get("run") or None, "updated_at": p.get("updated_at")})
    out = [{"name": k, "versions": sorted(v, key=lambda x: str(x.get("version")), reverse=True)}
           for k, v in grouped.items()]
    return {"project": project, "models": out, "count": sum(len(m["versions"]) for m in out), "degraded": err}


class PromoteRequest(BaseModel):
    project: str = "default"
    name: str
    version: str = "1"
    stage: str


@app.post("/api/studio/model/promote")
async def promote_model(req: PromoteRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Transition a model version's stage (staging → production …). A governed, attributable event: the model
    node is rewritten with the new stage, preserving its metrics/run lineage."""
    _require_write_token(authorization)
    if req.stage not in _MODEL_STAGES:
        raise HTTPException(status_code=422, detail=f"stage must be one of {_MODEL_STAGES}")
    coll = proj_collection(req.project)
    mid = _model_id(coll, req.name, req.version)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    existing = next((n for n in raw if n.get("id") == mid), None)
    if not existing:
        raise HTTPException(status_code=404, detail=f"model version not found: {mid}")
    p = dict(existing.get("properties") or {})
    prev = p.get("stage", "none")
    p["stage"] = req.stage
    p["promoted_at"] = _now_iso()
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": mid, "labels": existing.get("labels") or [coll, "Model"], "properties": p})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
    return {"model_id": mid, "name": req.name, "version": req.version,
            "stage": req.stage, "from_stage": prev, "proof_carrying": True}


# ── WS#48 (Wave 4): GraphRAG community detection. MEET Microsoft GraphRAG's core move — partition the knowledge
# graph into communities so global/local search can reason over community structure, not just isolated nodes. This
# is the DETERMINISTIC graph half (label propagation — no LLM, reproducible), the foundation community summaries
# attach to. BEAT: communities are detected over the PROOF-CARRYING graph, and each carries its epistemic profile
# (a quality signal), so a summary can be weighted by grounding — and the eventual summary text is FRONTIER-authored
# and attributed, never a local model. ─────────────────────────────────────────────────────────────────────────
def _detect_communities(nodes: list[dict[str, Any]], edges: list[dict[str, Any]],
                        max_iter: int = 50) -> dict[str, str]:
    """Louvain-style modularity optimisation (single level) → {node_id: community_seed}. Deterministic: nodes are
    swept in sorted order and ties broken by the smallest community id, so identical graphs yield identical
    communities. Robust to the 'monster community' collapse that plain label propagation suffers on bridged graphs
    — a node only joins a neighbour community when the modularity gain beats staying put."""
    ids = sorted({n.get("id") for n in nodes if n.get("id")})
    adj: dict[str, dict[str, float]] = {i: {} for i in ids}
    m = 0.0
    for e in edges:
        f, t = e.get("from"), e.get("to")
        if f not in adj or t not in adj or f == t:
            continue
        props = e.get("properties")
        w = float(props.get("weight", 1) or 1) if isinstance(props, dict) else 1.0
        adj[f][t] = adj[f].get(t, 0.0) + w
        adj[t][f] = adj[t].get(f, 0.0) + w
        m += w
    if m == 0:
        return {i: i for i in ids}
    two_m = 2.0 * m
    k = {i: sum(adj[i].values()) for i in ids}          # weighted degree
    comm = {i: i for i in ids}
    tot = {i: k[i] for i in ids}                        # Σ of degrees in each community
    for _ in range(max_iter):
        moved = False
        for i in ids:
            ci = comm[i]
            tot[ci] -= k[i]                             # take i out of its community
            neigh_w: dict[str, float] = {}
            for nb, w in adj[i].items():
                neigh_w[comm[nb]] = neigh_w.get(comm[nb], 0.0) + w
            best_c = ci
            best_gain = neigh_w.get(ci, 0.0) - k[i] * tot.get(ci, 0.0) / two_m
            for c, w_in in sorted(neigh_w.items()):     # deterministic order
                gain = w_in - k[i] * tot.get(c, 0.0) / two_m
                if gain > best_gain + 1e-12 or (abs(gain - best_gain) <= 1e-12 and c < best_c):
                    best_gain, best_c = gain, c
            comm[i] = best_c
            tot[best_c] = tot.get(best_c, 0.0) + k[i]
            if best_c != ci:
                moved = True
        if not moved:
            break
    # canonicalise each community's label to the smallest member id (stable, readable)
    canon: dict[str, str] = {}
    for i in ids:
        c = comm[i]
        canon[c] = i if c not in canon else min(canon[c], i)
    return {i: canon[comm[i]] for i in ids}


@app.get("/api/studio/communities")
async def communities(project: str = "default", min_size: int = 1,
                      _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Detect communities in the project knowledge graph (deterministic label propagation) — the GraphRAG
    foundation. Each community reports its size, top members by degree, and its epistemic distribution (quality),
    plus the inter-community link count. Summaries attach next and are frontier-authored + attributed."""
    coll = proj_collection(project)
    nodes, edges, err = await _fetch_graph(coll, 3000)
    by_id = {n.get("id"): n for n in nodes}
    degree: dict[str, int] = {}
    for e in edges:
        f, t = e.get("from"), e.get("to")
        if f in by_id:
            degree[f] = degree.get(f, 0) + 1
        if t in by_id:
            degree[t] = degree.get(t, 0) + 1
    label = _detect_communities(nodes, edges)
    groups: dict[str, list[str]] = {}
    for nid, lab in label.items():
        groups.setdefault(lab, []).append(nid)
    # inter-community edges (cross-links between communities) = the global-search backbone
    inter = sum(1 for e in edges
                if e.get("from") in label and e.get("to") in label
                and label[e["from"]] != label[e["to"]])
    out = []
    for lab, members in groups.items():
        if len(members) < max(1, min_size):
            continue
        epi: dict[str, int] = {}
        for m in members:
            mode = ((by_id.get(m) or {}).get("properties") or {}).get("epistemic_mode")
            if mode:
                epi[mode] = epi.get(mode, 0) + 1
        top = sorted(members, key=lambda m: (-degree.get(m, 0), m))[:6]
        top_named = [{"id": m, "label": ((by_id.get(m) or {}).get("properties") or {}).get("name") or m,
                      "degree": degree.get(m, 0)} for m in top]
        out.append({"community": f"{coll}:community:{hashlib.sha256(lab.encode()).hexdigest()[:10]}",
                    "seed": lab, "size": len(members), "top_members": top_named,
                    "epistemic_distribution": epi})
    out.sort(key=lambda c: c["size"], reverse=True)
    return {
        "project": project, "collection": coll,
        "communities": out, "count": len(out),
        "nodes": len(nodes), "inter_community_edges": inter,
        "algorithm": "louvain-modularity (deterministic, single-level)",
        "beat": "communities are detected over the proof-carrying graph — each carries its epistemic profile (grounding), and summaries are frontier-authored + attributed, never a local model",
        "degraded": err,
    }


# ── WS#31 (reframed): the pay-gated EXECUTION plane. The honest model — execution is a PROVISIONED, entitlement-
# gated full service (Databricks/Foundry don't give free compute either). The capability is declared and routable;
# the runtime only spins up when the project holds a paid compute entitlement. Multi-backend + sovereign: a small
# Spark namespace, a background Databricks connection, or self-hosted kind/k3s/k8s/DinD in the paid mesh. Fail-
# closed on BOTH the write token and the entitlement. BEAT: every run is a proof-carrying fact and emits a
# governed, replayable receipt — and you're not locked to one vendor's compute (unlike Foundry/Databricks). ─────
COMPUTE_BACKENDS = [
    {"id": "mesh-k8s", "kind": "sovereign", "default": True,
     "note": "self-hosted sandbox on the paid mesh (kind / k3s / k8s / DinD) — your hardware, our orchestration"},
    {"id": "spark", "kind": "sovereign",
     "note": "a small Spark namespace on the mesh — distributed dataframe/SQL compute"},
    {"id": "databricks", "kind": "external",
     "note": "connect to your own Databricks workspace in the background — bring-your-own compute"},
]
_BACKEND_IDS = {b["id"] for b in COMPUTE_BACKENDS}
EXEC_KINDS = {"notebook-cell", "pipeline-step", "job", "query"}
# Comma-sep entitlement tokens: "*" (all), "<project>" (any backend for it), "<project>:<backend>". Empty (default)
# = nothing entitled → every execute returns 402 (capability available, not provisioned). Pay-gating, fail-closed.
STUDIO_COMPUTE_ENTITLEMENTS = os.getenv("STUDIO_COMPUTE_ENTITLEMENTS", "")


def _entitlements() -> set[str]:
    return {t.strip() for t in STUDIO_COMPUTE_ENTITLEMENTS.split(",") if t.strip()}


def _compute_entitled(project: str, backend: str) -> bool:
    ents = _entitlements()
    return "*" in ents or project in ents or f"{project}:{backend}" in ents


@app.get("/api/studio/compute")
async def compute(project: str = "default",
                  _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The execution backends and this project's entitlement status. Every backend is AVAILABLE (declared +
    routable); it only runs when the project holds a paid compute entitlement — a full service, not by-the-drink."""
    backends = [{**b, "entitled": _compute_entitled(project, b["id"])} for b in COMPUTE_BACKENDS]
    return {"project": project, "backends": backends,
            "entitled_any": any(b["entitled"] for b in backends),
            "model": "pay-gated full service — capability available, runtime provisioned only when entitled (like Databricks/Foundry), but sovereign + multi-backend + proof-carrying receipts"}


class ExecuteRequest(BaseModel):
    project: str = "default"
    kind: str = "notebook-cell"            # notebook-cell | pipeline-step | job | query
    backend: str = "mesh-k8s"
    ref: str | None = None                 # what to run (notebook/cell/pipeline id)
    code: str | None = None                # inline payload — for backend='spark' this is the Spark SQL
    data: list[dict[str, Any]] = []        # inline rows for the spark job (registered as table `t`)
    note: str | None = None
    actor: str | None = None


async def _spark_submit(sql: str, data: list[dict[str, Any]], correlation: str) -> dict[str, Any] | None:
    """Dispatch a job to the spark-runner service. Returns its {rows, receipt} on success, else None (unset URL,
    unreachable, or a non-200 — the caller then degrades to the run-ledger). Isolated so execute() is testable."""
    if not SPARK_RUNNER_URL:
        return None
    try:
        async with httpx.AsyncClient(timeout=float(os.getenv("SPARK_TIMEOUT", "30"))) as sc:
            r = await sc.post(f"{SPARK_RUNNER_URL}/v1/submit",
                              json={"sql": sql, "data": data, "job_id": correlation},
                              headers={"authorization": f"Bearer {SPARK_RUNNER_TOKEN}"})
        return r.json() if r.status_code == 200 else None
    except Exception:  # noqa: BLE001 — spark-runner health is not our contract; degrade to the ledger
        return None


@app.post("/api/studio/execute")
async def execute(req: ExecuteRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Submit an execution to an entitled backend. Fail-closed twice: no write token → 503; no compute entitlement
    → 402 (capability available, not provisioned — provision a paid entitlement to run). When entitled, the run is
    recorded as a proof-carrying Execution fact and a governed, replayable receipt is emitted; the actual compute
    runs on the entitled backend's runtime (mesh / spark / databricks)."""
    _require_write_token(authorization)
    if req.backend not in _BACKEND_IDS:
        raise HTTPException(status_code=422, detail=f"backend must be one of {sorted(_BACKEND_IDS)}")
    if req.kind not in EXEC_KINDS:
        raise HTTPException(status_code=422, detail=f"kind must be one of {sorted(EXEC_KINDS)}")
    if not _compute_entitled(req.project, req.backend):
        raise HTTPException(status_code=402, detail={
            "status": "entitlement_required", "capability": "available", "backend": req.backend,
            "message": "compute is a paid, provisioned service — not spun up by the drink. Provision a compute "
                       "entitlement for this project/backend to run.",
            "backends": [b["id"] for b in COMPUTE_BACKENDS]})
    coll = proj_collection(req.project)
    payload_hash = hashlib.sha256((req.code or req.ref or "").encode()).hexdigest()
    correlation = f"exec-{payload_hash[:12]}"
    eid = f"{coll}:execution:{payload_hash[:12]}"
    prov = _workbench_prov(coll, "attested", req.actor or "studio/execute")
    prov["extractor"] = "lattice-studio/execution-v0"
    receipt = {"correlation_id": correlation, "service": "lattice-studio", "kind": req.kind,
               "backend": req.backend, "replayable": True, "payload_sha256": payload_hash,
               "bundle_ref": f"/v1/receipts/lattice-studio/{correlation}"}
    # backend='spark' → actually RUN the job on the sovereign spark-runner and chain its receipt; otherwise (or if
    # spark-runner is unreachable) the run is recorded in the governed ledger for the backend's runtime to pick up.
    status = "dispatched"
    rows = None
    spark_receipt = None
    if req.backend == "spark" and (req.code or "").strip():
        out = await _spark_submit(req.code, req.data, correlation)
        if out:
            status = "completed"
            rows = out.get("rows")
            spark_receipt = out.get("receipt")
            receipt["chained"] = spark_receipt.get("correlation_id") if isinstance(spark_receipt, dict) else None
    props = {"kind": req.kind, "backend": req.backend, "ref": req.ref or "", "status": status,
             "correlation_id": correlation, "receipt_bundle": receipt["bundle_ref"],
             "payload_sha256": payload_hash, "note": req.note or "", "submitted_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": eid, "labels": [coll, "Execution", "Run"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        if req.ref:                          # lineage: the execution ran a pipeline/notebook/etc
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                       json={"label": "executed", "from": eid, "to": req.ref, "properties": prov})
    return {"execution_id": eid, "backend": req.backend, "kind": req.kind, "status": status,
            "receipt": receipt, "spark_receipt": spark_receipt, "rows": rows, "proof_carrying": True,
            "note": (f"ran on the sovereign spark-runner; receipt chained" if status == "completed"
                     else f"run recorded + receipt emitted; compute executes on the entitled {req.backend} runtime")}


@app.get("/api/studio/executions")
async def executions(project: str = "default",
                     _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The execution ledger — every submitted run as a proof-carrying fact with its backend, status and receipt."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = []
    for n in raw:
        if "Execution" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        out.append({"execution_id": n.get("id"), "kind": p.get("kind"), "backend": p.get("backend"),
                    "status": p.get("status", "dispatched"), "ref": p.get("ref") or None,
                    "correlation_id": p.get("correlation_id"), "receipt_bundle": p.get("receipt_bundle"),
                    "submitted_at": p.get("submitted_at")})
    out.sort(key=lambda x: x.get("submitted_at") or "", reverse=True)
    return {"project": project, "executions": out, "count": len(out), "degraded": err}


# ── WS#50: GAIA stewardship writeback — the estate's canonical living-knowledge governance, wired to the moat.
# A steward's decision (keeper / successor / developmental phase / acknowledged abandonment signals) is persisted
# as gaia_* properties (the applyStewardship pattern), proof-carrying + receipted + reversible. THE alignment:
# GAIA invariant #2 ("model inference must not promote developmental state to canonical human-impacting truth")
# is enforced THROUGH epistemic status — a phase from a human steward is `attested`, from a model/agent `derived`.
# GAIA governance and the proof-carrying moat are the same mechanism. ────────────────────────────────────────────
class StewardRequest(BaseModel):
    project: str = "default"
    target: str                            # the node being stewarded (a GAIA LIVING_ENTITY / any node)
    keeper: str | None = None
    successor: str | None = None
    phase: str | None = None               # an OntogenesisState phase (seed…termination)
    resolve_signals: list[str] = []        # abandonment signals acknowledged/handled
    note: str | None = None
    actor: str | None = None
    actor_kind: str = "human"              # human | steward | keeper → attested; model | agent → derived


@app.post("/api/studio/steward")
async def steward(req: StewardRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Apply a GAIA steward decision to a node — governed, proof-carrying, receipted, reversible. Enforces the
    GAIA invariants: a developmental phase asserted by a model/agent is recorded as `derived`, not `attested`
    (invariant #2, via epistemic status). Fail-closed."""
    _require_write_token(authorization)
    if req.phase and req.phase not in gaia.ONTOGENESIS_PHASES:
        raise HTTPException(status_code=422, detail=f"phase must be one of {gaia.ONTOGENESIS_PHASES}")
    bad = [s for s in req.resolve_signals if s not in gaia.ABANDONMENT_SIGNALS]
    if bad:
        raise HTTPException(status_code=422, detail=f"unknown abandonment signals: {bad}")
    coll = proj_collection(req.project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    target_node = next((n for n in raw if n.get("id") == req.target), None)
    before = dict((target_node.get("properties") if target_node else {}) or {})
    labels = (target_node.get("labels") if target_node else None) or [coll, "LIVING_ENTITY"]

    # the writeback's epistemic status honors the invariant: a model steward-decision is derived, not canonical
    phase_epi = gaia.phase_epistemic(req.actor_kind)
    invariant = req.phase and phase_epi == "derived"
    prov = _workbench_prov(coll, "attested" if gaia.phase_epistemic(req.actor_kind) == "attested" else "derived",
                           req.actor or "studio/steward")
    prov["extractor"] = "lattice-studio/gaia-steward-v0"

    new_props = dict(before)
    changed: list[str] = []
    if req.keeper is not None:
        new_props["gaia_keeper"] = req.keeper; changed.append("keeper")
    if req.successor is not None:
        new_props["gaia_successor"] = req.successor; changed.append("successor")
    if req.phase:
        new_props["gaia_phase_override"] = req.phase
        new_props["gaia_phase_epistemic"] = phase_epi                 # invariant #2 in the data
        changed.append(f"phase={req.phase} ({phase_epi})")
    if req.resolve_signals:
        merged = sorted({*[s for s in str(before.get("gaia_resolved_signals", "")).split(",") if s], *req.resolve_signals})
        new_props["gaia_resolved_signals"] = ",".join(merged); changed.append("resolved_signals")
    if req.note is not None:
        new_props["gaia_steward_note"] = req.note
    new_props["gaia_reviewed_at"] = _now_iso()
    new_props.update(prov)

    dhash = hashlib.sha256((req.target + _now_iso() + ",".join(changed)).encode()).hexdigest()
    correlation = f"stw-{dhash[:12]}"
    dec_id = f"{coll}:stewardship:{dhash[:12]}"
    receipt = {"correlation_id": correlation, "service": "lattice-studio", "kind": "gaia-steward",
               "replayable": True, "bundle_ref": f"/v1/receipts/lattice-studio/{correlation}"}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": req.target, "labels": labels, "properties": new_props})
        if werr:
            raise HTTPException(status_code=502, detail=f"stewardship writeback failed: {werr}")
        dec_props = {"target": req.target, "changed": ", ".join(changed), "actor_kind": req.actor_kind,
                     "before_state": json.dumps(before, default=str), "correlation_id": correlation,
                     "receipt_bundle": receipt["bundle_ref"], "revoked": False, "at": _now_iso(), **prov}
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                   json={"id": dec_id, "labels": [coll, "STEWARDSHIP_RECORD", "Run"], "properties": dec_props})
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "STEWARD_OF", "from": dec_id, "to": req.target, "properties": prov})
    return {"stewardship_id": dec_id, "target": req.target, "state": gaia.stewardship_of(new_props),
            "changed": changed, "receipt": receipt, "reversible": True, "proof_carrying": True,
            "epistemic_mode": prov["epistemic_mode"],
            "invariant_applied": ("GAIA-2: model inference recorded as `derived`, not promoted to canonical truth"
                                  if invariant else None),
            "degraded": err}


@app.get("/api/studio/steward")
async def steward_state(project: str = "default", target: str = "",
                        _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """Read a node's persisted GAIA stewardship state — keeper, successor, developmental phase (+ its epistemic
    status), acknowledged abandonment signals. Also surfaces a light derived signal (orphaned if it has no edges)."""
    if not target:
        raise HTTPException(status_code=422, detail="target node id required")
    coll = proj_collection(project)
    nodes, edges, err = await _fetch_graph(coll, 3000)
    node = next((n for n in nodes if n.get("id") == target), None)
    state = gaia.stewardship_of(node.get("properties") if node else {})
    degree = sum(1 for e in edges if e.get("from") == target or e.get("to") == target)
    derived_signals = ["orphaned_artifact"] if degree == 0 else []
    derived_signals = [s for s in derived_signals if s not in state["resolved_signals"]]   # honor acknowledgements
    return {"project": project, "target": target, "stewardship": state,
            "derived_signals": derived_signals, "degree": degree,
            "phases": gaia.ONTOGENESIS_PHASES, "abandonment_signals": gaia.ABANDONMENT_SIGNALS,
            "invariants": gaia.GAIA_INVARIANTS, "degraded": err}


# ── WS#51: GAIA World Model — the decision-grade world-signal + the PROMOTION MEMBRANE as epistemic status. This
# is the three-twin unification made real: a gaia:WorldSignal (a governed Earth observation) is a proof-carrying
# HellGraph fact whose promotion state (EvidenceOnly → ReviewRequired → Promoted) IS its epistemic status. It is
# SHACL-validated against the real GAIA world-signals shapes on submit, and — invariant #2 — a model may propose
# but only a human/policy actor may Promote to canonical. Same discipline as the knowledge & human twins. ────────
class WorldSignalRequest(BaseModel):
    project: str = "default"
    feature_id: str
    signal_type: str = "feature_registry"     # feature_registry | foot_traffic | weather | concordance | …
    value: Any = None
    confidence: float | None = None
    geo_anchor: str | None = None             # a spatial-temporal anchor id/ref (GeoAnchor)
    evidence_refs: list[str] = []             # SourceEvidence ids — required to ever be Promoted
    actor: str | None = None
    actor_kind: str = "human"


def _ws_local(class_curie: str) -> str:
    return class_curie.split(":", 1)[-1]


@app.post("/api/studio/worldsignal")
async def submit_worldsignal(req: WorldSignalRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Submit a governed GAIA world-signal. It enters at PromotionState=EvidenceOnly (NOT canonical) — the writeback
    is SHACL-validated against the real GAIA shapes, and its epistemic status is derived from the promotion state
    (evidence-only → observed/derived by actor). Fail-closed."""
    _require_write_token(authorization)
    if req.signal_type not in gaia.SIGNAL_TYPES:
        raise HTTPException(status_code=422, detail=f"signal_type must be one of {sorted(gaia.SIGNAL_TYPES)}")
    fid = req.feature_id.strip()
    if not fid:
        raise HTTPException(status_code=422, detail="feature_id required")
    class_curie = gaia.SIGNAL_TYPES[req.signal_type]
    coll = proj_collection(req.project)
    state = "EvidenceOnly"                     # a fresh signal is always evidence-only, never born canonical
    epi = gaia.epistemic_for_promotion(state, req.actor_kind)
    sid = f"{coll}:worldsignal:{_norm(fid).replace(' ', '_')}"
    prov = _workbench_prov(coll, epi, req.actor or "gaia/worldsignal")
    prov["extractor"] = "lattice-studio/gaia-worldsignal-v0"
    gaia_props = {"gaia:hasFeatureId": fid, "gaia:hasPromotionState": state}
    if req.confidence is not None:
        gaia_props["gaia:hasConfidence"] = req.confidence
    conforms, violations = shacl.validate_gaia(class_curie, gaia_props)
    if not conforms:
        raise HTTPException(status_code=422, detail={"message": "world-signal rejected — not conformant to the GAIA shapes",
                                                     "class": class_curie, "violations": violations})
    props = {**gaia_props, "signal_type": req.signal_type, "gaia_class": class_curie,
             "promotion_state": state, "feature_id": fid, "value": req.value if req.value is not None else "",
             "confidence": req.confidence if req.confidence is not None else "",
             "geo_anchor": req.geo_anchor or "", "evidence_refs": json.dumps(req.evidence_refs),
             "actor_kind": req.actor_kind, "submitted_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": sid, "labels": [coll, "WorldSignal", _ws_local(class_curie)], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        for ev in req.evidence_refs:
            await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                       json={"label": "has_evidence", "from": sid, "to": ev, "properties": prov})
    return {"signal_id": sid, "class": class_curie, "promotion_state": state, "epistemic_mode": epi,
            "admissible_for_promotion": bool(req.evidence_refs), "proof_carrying": True,
            "note": "a world-signal is evidence-only until a policy decision Promotes it — its promotion state IS its epistemic status"}


class PromoteSignalRequest(BaseModel):
    project: str = "default"
    signal: str                               # world-signal node id
    to_state: str                             # ReviewRequired | Rejected | Promoted
    policy_id: str | None = None
    actor: str | None = None
    actor_kind: str = "human"


@app.post("/api/studio/worldsignal/promote")
async def promote_worldsignal(req: PromoteSignalRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """A GAIA PromotionDecision — move a world-signal across the promotion membrane, recording a replayable
    DecisionLedgerEntry. Enforces invariant #2 (a model/agent may NOT Promote to canonical) and admissibility (a
    signal cannot be Promoted without source evidence). Promotion updates the signal's epistemic status."""
    _require_write_token(authorization)
    if req.to_state not in ("ReviewRequired", "Rejected", "Promoted"):
        raise HTTPException(status_code=422, detail="to_state must be ReviewRequired | Rejected | Promoted")
    if not gaia.can_promote_to(req.actor_kind, req.to_state):
        raise HTTPException(status_code=403, detail={
            "message": "GAIA invariant #2 — model inference may not promote to canonical truth; only a human/policy actor Promotes",
            "actor_kind": req.actor_kind, "to_state": req.to_state})
    coll = proj_collection(req.project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    node = next((n for n in raw if n.get("id") == req.signal), None)
    if not node:
        raise HTTPException(status_code=404, detail=f"world-signal not found: {req.signal}")
    p = dict(node.get("properties") or {})
    try:
        evidence = json.loads(p.get("evidence_refs") or "[]")
    except (ValueError, TypeError):
        evidence = []
    if req.to_state == "Promoted" and not evidence:
        raise HTTPException(status_code=422, detail="not admissible — a world-signal cannot be Promoted without source evidence")
    prev = p.get("promotion_state", "EvidenceOnly")
    new_epi = gaia.epistemic_for_promotion(req.to_state, p.get("actor_kind", "human"))
    p["gaia:hasPromotionState"] = req.to_state
    p["promotion_state"] = req.to_state
    p["epistemic_mode"] = new_epi
    p["promoted_at"] = _now_iso()
    dhash = hashlib.sha256((req.signal + req.to_state + _now_iso()).encode()).hexdigest()
    correlation = f"dec-{dhash[:12]}"
    dec_id = f"{coll}:decision:{dhash[:12]}"
    prov = _workbench_prov(coll, "attested" if gaia.is_human(req.actor_kind) else "derived", req.actor or "gaia/policy")
    prov["extractor"] = "lattice-studio/gaia-promotion-v0"
    receipt = {"correlation_id": correlation, "service": "lattice-studio", "kind": "gaia-promotion",
               "replayable": True, "bundle_ref": f"/v1/receipts/lattice-studio/{correlation}"}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": req.signal, "labels": node.get("labels") or [coll, "WorldSignal"], "properties": p})
        if werr:
            raise HTTPException(status_code=502, detail=f"promotion writeback failed: {werr}")
        dec_props = {"gaia:hasDecisionId": correlation, "gaia:hasDecisionType": "feature_promotion",
                     "gaia:hasPolicyId": req.policy_id or "policy/default",
                     "from_state": prev, "to_state": req.to_state, "signal": req.signal,
                     "correlation_id": correlation, "receipt_bundle": receipt["bundle_ref"], "at": _now_iso(), **prov}
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                   json={"id": dec_id, "labels": [coll, "DecisionLedgerEntry", "PromotionDecision", "Run"], "properties": dec_props})
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "promotesSignal", "from": dec_id, "to": req.signal, "properties": prov})
    return {"decision_id": dec_id, "signal": req.signal, "from_state": prev, "to_state": req.to_state,
            "epistemic_mode": new_epi, "receipt": receipt, "proof_carrying": True,
            "canonical": req.to_state == "Promoted"}


@app.get("/api/studio/worldsignals")
async def worldsignals(project: str = "default", state: str = "",
                       _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The project's GAIA world-signals with their promotion state + epistemic status + evidence."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = []
    for n in raw:
        if "WorldSignal" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        if state and p.get("promotion_state") != state:
            continue
        try:
            evidence = json.loads(p.get("evidence_refs") or "[]")
        except (ValueError, TypeError):
            evidence = []
        out.append({"signal_id": n.get("id"), "feature_id": p.get("feature_id"), "signal_type": p.get("signal_type"),
                    "promotion_state": p.get("promotion_state", "EvidenceOnly"), "epistemic_mode": p.get("epistemic_mode"),
                    "canonical": p.get("promotion_state") == "Promoted", "evidence_count": len(evidence),
                    "confidence": p.get("confidence") or None, "submitted_at": p.get("submitted_at")})
    out.sort(key=lambda x: x.get("submitted_at") or "", reverse=True)
    return {"project": project, "world_signals": out, "count": len(out), "degraded": err}


@app.get("/api/studio/gaia/ontology")
async def gaia_ontology(_auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The GAIA world-signals promotion membrane — and the mapping that unifies the estate: a promotion state IS
    an epistemic status. The same governed-evidence discipline runs across all three digital twins."""
    membrane = [{"state": s, "epistemic_human": gaia.epistemic_for_promotion(s, "human"),
                 "epistemic_model": gaia.epistemic_for_promotion(s, "model"),
                 "canonical": s == "Promoted"} for s in gaia.PROMOTION_STATES]
    return {
        "ontology": "GAIA World Model — decision-grade world-signals", "namespace": gaia.GAIA_NS,
        "signal_types": gaia.SIGNAL_TYPES, "promotion_states": gaia.PROMOTION_STATES,
        "promotion_epistemic_membrane": membrane,
        "invariant": "GAIA-2: a model may propose (EvidenceOnly/ReviewRequired) but only a human/policy actor may Promote to canonical (attested)",
        "three_twins": {
            "knowledge": "HellGraph fact · epistemic status observed→attested",
            "human": "HDT Observation → OmegaState ABSENT→DELIVERED",
            "earth": "GAIA WorldSignal → PromotionState EvidenceOnly→Promoted",
            "note": "one governed-evidence discipline; the promotion membrane and the epistemic ladder are the same mechanism",
        },
    }


# ── WS#52: HDT — the Human Digital Twin, the THIRD twin (closes the triangle). An hdt:Observation carries an
# hdt:OmegaState (ABSENT→SEEDED→NORMALIZED→LINKED→TRUSTED→ACTIONABLE→DELIVERED) — the human twin's promotion
# lattice, which IS its epistemic status. Same discipline, same invariant: a model may seed/normalize but only a
# human/clinician/policy actor may DELIVER an observation to canonical, human-actionable truth. ─────────────────
class HdtObservationRequest(BaseModel):
    project: str = "default"
    subject: str                              # the person/twin the observation is about
    code: str                                 # the observation code (e.g. a FHIR/LOINC code)
    value: Any = None
    m_cbd: float | None = None                # KFS membership — Cognition
    m_cgt: float | None = None                # Values
    m_nhy: float | None = None                # Action
    actor: str | None = None
    actor_kind: str = "human"


@app.post("/api/studio/hdt/observation")
async def hdt_observation(req: HdtObservationRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Record a governed HDT observation about a person. It enters at OmegaState=SEEDED (not canonical); its
    epistemic status is derived from the OmegaState, and it carries the KFS membership triad. Proof-carrying,
    fail-closed."""
    _require_write_token(authorization)
    code = req.code.strip()
    if not code:
        raise HTTPException(status_code=422, detail="code required")
    coll = proj_collection(req.project)
    state = "SEEDED"
    epi = hdt.epistemic_for_omega(state, req.actor_kind)
    oid = f"{coll}:hdtobs:{hashlib.sha256((req.subject + code + _now_iso()).encode()).hexdigest()[:12]}"
    prov = _workbench_prov(coll, epi, req.actor or "hdt/observation")
    prov["extractor"] = "lattice-studio/hdt-observation-v0"
    props = {"subject": req.subject, "code": code, "value": req.value if req.value is not None else "",
             "omega_state": state, "hdt:hasOmegaState": state,
             "m_cbd": req.m_cbd if req.m_cbd is not None else "", "m_cgt": req.m_cgt if req.m_cgt is not None else "",
             "m_nhy": req.m_nhy if req.m_nhy is not None else "", "actor_kind": req.actor_kind,
             "recorded_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": oid, "labels": [coll, "HdtObservation", "Observation"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "about_subject", "from": oid, "to": req.subject, "properties": prov})
    return {"observation_id": oid, "subject": req.subject, "omega_state": state, "epistemic_mode": epi,
            "proof_carrying": True,
            "note": "an observation is SEEDED, not canonical — its OmegaState IS its epistemic status; only a human/policy DELIVERS it"}


class HdtPromoteRequest(BaseModel):
    project: str = "default"
    observation: str
    to_state: str                             # a later OmegaState (NORMALIZED … DELIVERED)
    actor: str | None = None
    actor_kind: str = "human"


@app.post("/api/studio/hdt/promote")
async def hdt_promote(req: HdtPromoteRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """An HDT EvaluationEvent — advance an observation along the OmegaState lattice, recording a replayable
    promotion event. Enforces the invariant: only a human/clinician/policy actor may DELIVER to canonical."""
    _require_write_token(authorization)
    if req.to_state not in hdt.OMEGA_STATES:
        raise HTTPException(status_code=422, detail=f"to_state must be one of {hdt.OMEGA_STATES}")
    if not hdt.can_promote_omega(req.actor_kind, req.to_state):
        raise HTTPException(status_code=403, detail={
            "message": "HDT invariant — a model may advance an observation but only a human/clinician/policy actor may DELIVER to canonical truth",
            "actor_kind": req.actor_kind, "to_state": req.to_state})
    coll = proj_collection(req.project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    node = next((n for n in raw if n.get("id") == req.observation), None)
    if not node:
        raise HTTPException(status_code=404, detail=f"observation not found: {req.observation}")
    p = dict(node.get("properties") or {})
    prev = p.get("omega_state", "ABSENT")
    new_epi = hdt.epistemic_for_omega(req.to_state, p.get("actor_kind", "human"))
    p["omega_state"] = req.to_state
    p["hdt:hasOmegaState"] = req.to_state
    p["epistemic_mode"] = new_epi
    p["promoted_at"] = _now_iso()
    ehash = hashlib.sha256((req.observation + req.to_state + _now_iso()).encode()).hexdigest()
    correlation = f"eval-{ehash[:12]}"
    ev_id = f"{coll}:evaluation:{ehash[:12]}"
    prov = _workbench_prov(coll, "attested" if hdt.is_human(req.actor_kind) else "derived", req.actor or "hdt/evaluation")
    prov["extractor"] = "lattice-studio/hdt-evaluation-v0"
    receipt = {"correlation_id": correlation, "service": "lattice-studio", "kind": "hdt-evaluation",
               "replayable": True, "bundle_ref": f"/v1/receipts/lattice-studio/{correlation}"}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": req.observation, "labels": node.get("labels") or [coll, "HdtObservation"], "properties": p})
        if werr:
            raise HTTPException(status_code=502, detail=f"promotion writeback failed: {werr}")
        ev_props = {"hdt:promotedFromState": prev, "hdt:promotedToState": req.to_state, "observation": req.observation,
                    "correlation_id": correlation, "receipt_bundle": receipt["bundle_ref"], "at": _now_iso(), **prov}
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                   json={"id": ev_id, "labels": [coll, "EvaluationEvent", "Run"], "properties": ev_props})
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "evaluates", "from": ev_id, "to": req.observation, "properties": prov})
    return {"evaluation_id": ev_id, "observation": req.observation, "from_state": prev, "to_state": req.to_state,
            "epistemic_mode": new_epi, "receipt": receipt, "proof_carrying": True,
            "canonical": req.to_state == "DELIVERED"}


@app.get("/api/studio/hdt")
async def hdt_observations(project: str = "default", subject: str = "",
                           _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The HDT observations (optionally for one subject) with their OmegaState + epistemic status + KFS membership."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = []
    for n in raw:
        if "HdtObservation" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        if subject and p.get("subject") != subject:
            continue
        out.append({"observation_id": n.get("id"), "subject": p.get("subject"), "code": p.get("code"),
                    "omega_state": p.get("omega_state", "ABSENT"), "epistemic_mode": p.get("epistemic_mode"),
                    "canonical": p.get("omega_state") == "DELIVERED", "recorded_at": p.get("recorded_at")})
    out.sort(key=lambda x: x.get("recorded_at") or "", reverse=True)
    return {"project": project, "observations": out, "count": len(out), "degraded": err}


@app.get("/api/studio/hdt/ontology")
async def hdt_ontology(_auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The HDT OmegaState lattice — and, with the other two, the CLOSURE: three digital twins, one discipline."""
    lattice = [{"state": s, "epistemic_human": hdt.epistemic_for_omega(s, "human"),
                "epistemic_model": hdt.epistemic_for_omega(s, "model"), "canonical": s == "DELIVERED"}
               for s in hdt.OMEGA_STATES]
    return {
        "ontology": "HDT — Human Digital Twin", "namespace": hdt.HDT_NS,
        "omega_states": hdt.OMEGA_STATES, "omega_epistemic_lattice": lattice, "kfs_triad": hdt.TRIAD_ROLES,
        "invariant": "a model may advance an observation but only a human/clinician/policy actor may DELIVER to canonical human-actionable truth",
        "three_twins_closed": {
            "knowledge": "HellGraph fact · observed→attested",
            "human": "HDT Observation · OmegaState ABSENT→DELIVERED",
            "earth": "GAIA WorldSignal · PromotionState EvidenceOnly→Promoted",
            "note": "all three twins now run one governed-evidence discipline in the same substrate — the promotion membrane IS the epistemic ladder, everywhere",
        },
    }


# ── WS#49: ONTOLOGY ACTIONS + writeback — the Foundry crown jewel (Workshop apps are built on Actions). A typed
# action defines a governed edit on an object type (set a property, set status, add a relation), and invoking it
# writes back to the graph. BEAT Foundry on every axis: their edits are bare edits; ours are PROOF-CARRYING
# (epistemic status + provenance on every writeback), RECEIPTED (a replayable receipt per invocation), REVERSIBLE
# (the before-state is snapshotted, so revoke restores it — governed undo as graph facts), and AGENT-INVOKABLE
# (each action publishes a machine-readable schema, so an agent can discover + invoke it — not UI-bound). ────────
class ActionParam(BaseModel):
    name: str
    type: str = "string"
    required: bool = True


class ActionEffect(BaseModel):
    op: str                                # set_property | set_status | add_edge
    property: str | None = None            # for set_property
    label: str | None = None               # for add_edge (the relation)
    value: Any = None                      # a literal value…
    value_from: str | None = None          # …or take it from an invocation arg


class ActionRequest(BaseModel):
    project: str = "default"
    name: str
    target_type: str                       # the object type (graph label) this action applies to
    params: list[ActionParam] = []
    effects: list[ActionEffect] = []
    description: str | None = None


def _action_id(coll: str, name: str) -> str:
    return f"{coll}:action:{_norm(name).replace(' ', '_')}"


@app.post("/api/studio/action")
async def define_action(req: ActionRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Define a typed ontology action (governed writeback) — the Foundry-Workshop primitive. Persisted as a
    proof-carrying Action node with a machine-readable schema, so both a Workshop-style UI and an AGENT can invoke
    it. Idempotent per (project, name)."""
    _require_write_token(authorization)
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="name required")
    if not req.effects:
        raise HTTPException(status_code=422, detail="at least one effect required")
    for e in req.effects:
        if e.op not in {"set_property", "set_status", "add_edge"}:
            raise HTTPException(status_code=422, detail=f"unsupported effect op: {e.op}")
    # TYPE the action against the REAL Ontogenesis ontology: target_type must be a class, and every effect's
    # property/relation must be declared on it (or inherited). This is what makes it an *ontology* action.
    resolved, errors = ontology.validate_action(req.target_type, [e.model_dump() for e in req.effects])
    if errors:
        raise HTTPException(status_code=422, detail={"message": "action does not conform to the ontology",
                                                     "target_type": req.target_type, "violations": errors})
    coll = proj_collection(req.project)
    aid = _action_id(coll, name)
    prov = _workbench_prov(coll, "attested", "studio/action")
    prov["extractor"] = "lattice-studio/action-v0"
    props = {"name": name, "target_type": req.target_type, "class_iri": resolved["iri"] if resolved else "",
             "params": json.dumps([p.model_dump() for p in req.params]),
             "effects": json.dumps([e.model_dump() for e in req.effects]),
             "description": req.description or "", "updated_at": _now_iso(), **prov}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": aid, "labels": [coll, "Action"], "properties": props})
        if werr:
            raise HTTPException(status_code=502, detail=f"graph write failed: {werr}")
    return {"action_id": aid, "name": name, "target_type": req.target_type,
            "class_iri": resolved["iri"] if resolved else None,
            "effects": len(req.effects), "proof_carrying": True, "agent_invokable": True, "ontology_typed": True}


@app.get("/api/studio/ontology")
async def ontology_classes(search: str = "", cls: str = "", limit: int = 40,
                           _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The REAL Ontogenesis ontology (817 classes / 621 relations, parsed from the ~/dev/ontogenesis TTL corpus and
    vendored) — replaces the old 2-item mock. `cls=<curie|label>` returns one class with its full (inherited)
    property set; `search=` filters the class list; otherwise returns the upper classes + a sample. This is the
    schema ontology actions are typed against."""
    if cls:
        c = ontology.resolve_class(cls)
        if not c:
            raise HTTPException(status_code=404, detail=f"class not found: {cls}")
        props = ontology.class_properties(c["iri"])
        return {"class": {**c, "inherited_properties": sorted(props.values(), key=lambda p: p["iri"])},
                "base_iri": ontology.base_iri()}
    all_classes = ontology.all_classes()
    if search:
        s = search.lower()
        hits = [c for c in all_classes if s in c["iri"].lower() or s in str(c.get("label", "")).lower()]
    else:
        hits = [c for c in all_classes if c["iri"].startswith("upper:")]     # the upper ontology by default
    slim = [{"iri": c["iri"], "label": c.get("label"), "subClassOf": c.get("subClassOf", []),
             "property_count": len(c.get("properties", []))} for c in hits[:max(1, min(limit, 200))]]
    return {"base_iri": ontology.base_iri(), "counts": ontology.counts(),
            "classes": slim, "returned": len(slim), "total_matched": len(hits),
            "note": "the real Ontogenesis OWL corpus; pass ?cls=<curie|label> for a class's properties, ?search= to filter"}


def _action_view(n: dict[str, Any]) -> dict[str, Any]:
    p = n.get("properties") or {}
    try:
        params = json.loads(p.get("params") or "[]")
    except (ValueError, TypeError):
        params = []
    try:
        effects = json.loads(p.get("effects") or "[]")
    except (ValueError, TypeError):
        effects = []
    return {"action_id": n.get("id"), "name": p.get("name"), "target_type": p.get("target_type"),
            "description": p.get("description") or None, "params": params, "effects": effects}


@app.get("/api/studio/actions")
async def actions(project: str = "default",
                  _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The project's ontology actions, each with its machine-readable schema (params + effects) — the discovery
    surface a Workshop UI *or* an agent uses to invoke them."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = [_action_view(n) for n in raw if "Action" in (n.get("labels") or [])]
    out.sort(key=lambda a: a.get("name") or "")
    return {"project": project, "actions": out, "count": len(out),
            "agent_note": "each action's params+effects are a machine-readable contract — agents discover and invoke, not just humans",
            "degraded": err}


class InvokeRequest(BaseModel):
    project: str = "default"
    action: str                            # action name
    target: str                            # the target node id the action acts on
    args: dict[str, Any] = {}
    actor: str | None = None


def _resolve(effect: dict[str, Any], args: dict[str, Any]) -> Any:
    return args.get(effect["value_from"]) if effect.get("value_from") else effect.get("value")


@app.post("/api/studio/action/invoke")
async def invoke_action(req: InvokeRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Invoke an action on a target — a governed writeback. Applies the effects to the target (proof-carrying,
    epistemic), snapshots the before-state for reversibility, records an ActionInvocation (the audit trail), and
    emits a replayable receipt. Fail-closed."""
    _require_write_token(authorization)
    coll = proj_collection(req.project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    if err:
        raise HTTPException(status_code=502, detail=f"graph read failed: {err}")
    aid = _action_id(coll, req.action)
    action_node = next((n for n in raw if n.get("id") == aid), None)
    if not action_node:
        raise HTTPException(status_code=404, detail=f"action not found: {req.action}")
    view = _action_view(action_node)
    missing = [p["name"] for p in view["params"] if p.get("required", True) and p["name"] not in req.args]
    if missing:
        raise HTTPException(status_code=422, detail=f"missing required args: {missing}")
    target_node = next((n for n in raw if n.get("id") == req.target), None)
    before = dict((target_node.get("properties") if target_node else {}) or {})
    labels = (target_node.get("labels") if target_node else None) or [coll, view["target_type"]]

    prov = _workbench_prov(coll, "attested", req.actor or "studio/action-invoke")
    prov["extractor"] = "lattice-studio/action-invoke-v0"
    new_props = dict(before)
    edge_effects: list[dict[str, Any]] = []
    applied: list[str] = []
    for e in view["effects"]:
        if e["op"] == "set_property":
            new_props[e["property"]] = _resolve(e, req.args)
            applied.append(f"set {e['property']}")
        elif e["op"] == "set_status":
            new_props["status"] = _resolve(e, req.args)
            applied.append("set status")
        elif e["op"] == "add_edge":
            edge_effects.append({"label": e.get("label") or "relates_to", "to": _resolve(e, req.args)})
            applied.append(f"+edge {e.get('label')}")
    new_props.update(prov)

    # GOVERNED writeback: validate the resulting node state against the real Ontogenesis SHACL shapes BEFORE it
    # commits. Non-conformant → 422 with the violations (Foundry's "submission criteria", but schema-real + proof-
    # carrying). No shape targeting the class → conforms. Degrades open if the validator is unavailable.
    class_iri = (action_node.get("properties") or {}).get("class_iri") or ""
    if class_iri:
        conforms, violations = shacl.validate_writeback(class_iri, new_props)
        if not conforms:
            raise HTTPException(status_code=422, detail={
                "message": "writeback rejected — does not conform to the ontology (SHACL)",
                "class": class_iri, "violations": violations})

    inv_hash = hashlib.sha256((aid + req.target + json.dumps(req.args, sort_keys=True, default=str) + _now_iso()).encode()).hexdigest()
    correlation = f"act-{inv_hash[:12]}"
    inv_id = f"{coll}:invocation:{inv_hash[:12]}"
    receipt = {"correlation_id": correlation, "service": "lattice-studio", "action": req.action,
               "replayable": True, "bundle_ref": f"/v1/receipts/lattice-studio/{correlation}"}
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                             json={"id": req.target, "labels": labels, "properties": new_props})   # the writeback
        if werr:
            raise HTTPException(status_code=502, detail=f"writeback failed: {werr}")
        for ee in edge_effects:
            if ee["to"]:
                await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                           json={"label": ee["label"], "from": req.target, "to": ee["to"], "properties": prov})
        inv_props = {"action": req.action, "target": req.target, "args": json.dumps(req.args, default=str),
                     "before_state": json.dumps(before, default=str), "applied": ", ".join(applied),
                     "correlation_id": correlation, "receipt_bundle": receipt["bundle_ref"],
                     "revoked": False, "invoked_at": _now_iso(), **prov}
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                   json={"id": inv_id, "labels": [coll, "ActionInvocation", "Run"], "properties": inv_props})
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "invoked", "from": inv_id, "to": aid, "properties": prov})
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                   json={"label": "acted_on", "from": inv_id, "to": req.target, "properties": prov})
    return {"invocation_id": inv_id, "action": req.action, "target": req.target, "applied": applied,
            "receipt": receipt, "reversible": True, "proof_carrying": True}


@app.get("/api/studio/action/invocations")
async def invocations(project: str = "default", target: str = "",
                      _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """The action audit trail — every invocation as a proof-carrying fact (action, target, effects, receipt,
    revoked?). Foundry has an edit history; ours is a queryable, receipted, reversible graph ledger."""
    coll = proj_collection(project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    out = []
    for n in raw:
        if "ActionInvocation" not in (n.get("labels") or []):
            continue
        p = n.get("properties") or {}
        if target and p.get("target") != target:
            continue
        out.append({"invocation_id": n.get("id"), "action": p.get("action"), "target": p.get("target"),
                    "applied": p.get("applied"), "revoked": bool(p.get("revoked")),
                    "correlation_id": p.get("correlation_id"), "receipt_bundle": p.get("receipt_bundle"),
                    "invoked_at": p.get("invoked_at")})
    out.sort(key=lambda x: x.get("invoked_at") or "", reverse=True)
    return {"project": project, "invocations": out, "count": len(out), "degraded": err}


class RevokeActionRequest(BaseModel):
    project: str = "default"
    invocation: str                        # the invocation node id


@app.post("/api/studio/action/revoke")
async def revoke_action(req: RevokeActionRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Governed undo — restore the target to the before-state the invocation snapshotted, and mark the invocation
    revoked. The moat Foundry lacks: a proof-carrying, reversible action ledger (property effects are restored;
    added relations are additive)."""
    _require_write_token(authorization)
    coll = proj_collection(req.project)
    raw, err = await _fetch_raw_nodes(coll, 2000)
    inv = next((n for n in raw if n.get("id") == req.invocation), None)
    if not inv:
        raise HTTPException(status_code=404, detail=f"invocation not found: {req.invocation}")
    p = inv.get("properties") or {}
    if p.get("revoked"):
        return {"invocation_id": req.invocation, "already_revoked": True}
    target = p.get("target")
    try:
        before = json.loads(p.get("before_state") or "{}")
    except (ValueError, TypeError):
        before = {}
    target_node = next((n for n in raw if n.get("id") == target), None)
    labels = (target_node.get("labels") if target_node else None) or [coll]
    prov = _workbench_prov(coll, "attested", "studio/action-revoke")
    prov["extractor"] = "lattice-studio/action-revoke-v0"
    restored = dict(before)
    restored.update(prov)
    new_inv = dict(p)
    new_inv["revoked"] = True
    new_inv["revoked_at"] = _now_iso()
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        if target_node:
            _, werr = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                                 json={"id": target, "labels": labels, "properties": restored})
            if werr:
                raise HTTPException(status_code=502, detail=f"restore failed: {werr}")
        await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                   json={"id": req.invocation, "labels": inv.get("labels") or [coll, "ActionInvocation", "Run"],
                         "properties": new_inv})
    return {"invocation_id": req.invocation, "target": target, "revoked": True,
            "restored_state": True, "proof_carrying": True}


# ── KE-1: the real extraction → HellGraph loop (proof-carrying, project-scoped) ──────────────────────────────────
# Deterministic entity + co-occurrence extraction (honest: not LLM). Every fact is written as a HellGraph atom with
# epistemic_mode="observed" + source provenance + the project label — the moat made real: not "entity X", but
# "entity X, observed, from source S, in project P, by extractor E". Neo4j stores the node; we store the node with
# its epistemic status and provenance, in one governed project scope.

# KE-1.1: stop-words (pronouns, determiners, conjunctions, common sentence-openers) that a capitalized-phrase
# extractor otherwise mistakes for entities. LEADING determiners are also stripped from phrases ("The Lattice
# Studio" → "Lattice Studio") so the entity is the noun, not the article.
_STOP = {
    "The", "This", "That", "These", "Those", "There", "Then", "When", "Where", "What", "Which", "While",
    "And", "But", "For", "With", "From", "Into", "It", "He", "She", "They", "We", "You", "I", "A", "An",
    "As", "At", "By", "Of", "On", "Or", "So", "To", "Its", "Our", "Their", "His", "Her", "Your", "My",
    "If", "In", "Is", "Are", "Was", "Were", "Be", "Has", "Have", "Had", "Not", "No", "Yes", "Also", "However",
}
_LEADING_DET = re.compile(r"^(The|This|That|These|Those|A|An|Its|Our|Their|His|Her|Your|My)\s+")
_PROPER = re.compile(r"\b([A-Z][a-zA-Z0-9]+(?:\s+[A-Z][a-zA-Z0-9]+)*)\b")
_ACRONYM = re.compile(r"\b([A-Z]{2,}[A-Za-z0-9]*)\b")


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())


def extract_facts(text: str, limit: int = 60) -> tuple[list[str], list[tuple[str, str]]]:
    """Deterministic entities (proper-noun phrases + acronyms) + co-occurrence relations (same sentence)."""
    entities: list[str] = []
    seen: set[str] = set()
    sentences = re.split(r"(?<=[.!?])\s+", text)
    relations: list[tuple[str, str]] = []
    rel_seen: set[tuple[str, str]] = set()
    for sent in sentences:
        local: list[str] = []
        for m in (*_PROPER.finditer(sent), *_ACRONYM.finditer(sent)):
            name = _LEADING_DET.sub("", m.group(1).strip()).strip()  # strip leading determiner
            if not name or name in _STOP or len(name) < 2:
                continue
            key = _norm(name)
            if key not in seen and len(entities) < limit:
                seen.add(key); entities.append(name)
            if key not in {_norm(x) for x in local}:
                local.append(name)
        # co-occurrence edges within a sentence (undirected, deduped)
        for i in range(len(local)):
            for j in range(i + 1, len(local)):
                a, b = sorted((_norm(local[i]), _norm(local[j])))
                if (a, b) not in rel_seen:
                    rel_seen.add((a, b)); relations.append((local[i], local[j]))
    return entities, relations[: limit * 2]


def _require_write_token(authorization: str) -> None:
    """Shared WRITE gate for every endpoint that mutates the shared HellGraph (extract + the node/edge
    workbench). Fail-CLOSED: if STUDIO_WRITE_TOKEN is unset, writes are refused (reads stay open), so a
    public ingress can never accept anonymous graph writes. Token provisioned out-of-band (Secret)."""
    if not STUDIO_WRITE_TOKEN:
        raise HTTPException(status_code=503, detail="studio writes disabled: STUDIO_WRITE_TOKEN unset (fail-closed)")
    if authorization.removeprefix("Bearer ").strip() != STUDIO_WRITE_TOKEN:
        raise HTTPException(status_code=401, detail="invalid write token")


def _ent_id(coll: str, name: str) -> str:
    """Project-scoped entity id — must match /extract's scheme so a hand-authored node and an extracted one
    with the same name are the SAME node (the workbench edits the same graph, not a parallel one)."""
    return f"{coll}:ent:{_norm(name).replace(' ', '_')}"


class ExtractRequest(BaseModel):
    project: str = "default"
    text: str
    source: str | None = None


@app.post("/api/studio/extract")
async def extract(req: ExtractRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    _require_write_token(authorization)
    coll = proj_collection(req.project)
    src = req.source or ("text:" + hashlib.sha256(req.text.encode()).hexdigest()[:16])
    entities, relations = extract_facts(req.text)
    ent_id = lambda name: _ent_id(coll, name)  # noqa: E731
    # KKO (KBpedia Knowledge Ontology) is the estate's UPPER ONTOLOGY. Extracted named entities are Peircean
    # Particulars (Secondness). epistemic_mode="observed" is itself Peircean (KKO formalizes the inference
    # trichotomy induced/deduced/abduced as kko:Methodeutic) — so the provenance is standards-grounded, not ad-hoc.
    prov = {"epistemic_mode": "observed", "source": src, "extractor": "lattice-studio/deterministic-v0",
            "project": coll, "kko_type": "Particulars"}

    written_nodes = written_edges = 0
    errors: list[str] = []
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        # nodes first (concurrently), then edges (concurrently)
        node_calls = [
            _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                 json={"id": ent_id(n), "labels": [coll, "Entity"], "properties": {"name": n, **prov}})
            for n in entities
        ]
        for _, err in await asyncio.gather(*node_calls):
            if err: errors.append(err)
            else: written_nodes += 1
        edge_calls = [
            _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                 json={"label": "co_occurs", "from": ent_id(a), "to": ent_id(b), "properties": prov})
            for a, b in relations
        ]
        for _, err in await asyncio.gather(*edge_calls):
            if err: errors.append(err)
            else: written_edges += 1

    return {
        "project": req.project, "projectCollection": coll, "source": src,
        "extracted": {"entities": len(entities), "relations": len(relations)},
        "written": {"nodes": written_nodes, "edges": written_edges},
        "provenance": prov,
        "sample": entities[:10],
        "errors": errors[:5] or None,
    }


def _chunk_text(text: str, max_chars: int = 4000) -> list[str]:
    """Paragraph-preserving chunks for the IE hop (ie-engine caps entities/relations PER CALL, so chunking
    is what lets a full document through, not just its first page). Oversized paragraphs split on sentence
    boundaries; a pathological single sentence is hard-split rather than dropped."""
    chunks: list[str] = []
    buf = ""
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        if len(buf) + len(para) + 2 <= max_chars:
            buf = f"{buf}\n\n{para}" if buf else para
            continue
        if buf:
            chunks.append(buf); buf = ""
        if len(para) <= max_chars:
            buf = para
            continue
        for sent in re.split(r"(?<=[.!?])\s+", para):
            while len(sent) > max_chars:                     # pathological unbroken run
                chunks.append(sent[:max_chars]); sent = sent[max_chars:]
            if len(buf) + len(sent) + 1 > max_chars:
                chunks.append(buf); buf = sent
            else:
                buf = f"{buf} {sent}" if buf else sent
    if buf:
        chunks.append(buf)
    return chunks


class IngestDocumentRequest(BaseModel):
    project: str = "default"
    text: str
    filename: str | None = None
    source: str | None = None


@app.post("/api/studio/ingest-document")
async def ingest_document(req: IngestDocumentRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """The document→linked-knowledge pipeline (ST024): chunk → ie-engine (real spaCy NER + SVO relations)
    → entity-resolution (proof-carrying golden records) → hellgraph upsert under ONE canonical id scheme.

    Before this endpoint the estate had extraction, ER, and the graph as disconnected hops with two
    incompatible id schemes (ie:<slug> vs <coll>:ent:<slug>) and NOBODY calling ER — the same real-world
    entity landed as different nodes per path. Here every mention is resolved to a golden record first and
    written as _ent_id(coll, canonical_name), the SAME scheme /extract and the workbench use, so document-
    extracted, hand-authored, and regex-extracted facts about one entity converge on one node.

    Fail-soft at each hop: IE unreachable → deterministic extract_facts fallback (tagged in provenance);
    ER unreachable → identity resolution (each mention its own entity, error surfaced). Graph writes carry
    the full studio provenance stamp + doc_sha so every fact traces back to the source document."""
    _require_write_token(authorization)
    text = req.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="empty document")
    return await _run_ingest_pipeline(req.project, text, req.filename, req.source)


async def _run_ingest_pipeline(project: str, text: str, filename: str | None, source: str | None,
                               extra_prov: dict[str, Any] | None = None) -> dict[str, Any]:
    """The shared IE→ER→graph pipeline behind /ingest-document (raw text) and /ingest-file (converted files).
    Caller is responsible for the write gate and for handing in non-empty text."""
    coll = proj_collection(project)
    doc_sha = hashlib.sha256(text.encode()).hexdigest()
    src = source or f"doc:{doc_sha[:16]}"
    chunks = _chunk_text(text)

    errors: list[str] = []
    mentions: dict[str, str] = {}          # mention text → entity type ("" when unknown)
    relations: list[dict[str, str]] = []   # {from, relation, to} in mention-text space
    claims: list[dict[str, Any]] = []
    extractor = "lattice-studio/ie-pipeline-v1"

    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        # 1) EXTRACT — ie-engine per chunk, concurrently. Topics are ambience, not entities → skipped.
        ie_results = await asyncio.gather(
            *[_req(client, "POST", f"{IE_ENGINE_URL}/extract", json={"text": c}) for c in chunks]
        )
        ie_ok = False
        for body, err in ie_results:
            if err or not isinstance(body, dict):
                if err:
                    errors.append(f"ie-engine: {err}")
                continue
            ie_ok = True
            for e in body.get("entities", []):
                name, typ = str(e.get("text", "")).strip(), str(e.get("type", ""))
                if name and typ != "Topic":
                    mentions.setdefault(name, typ)
            relations.extend(
                {"from": str(r.get("from", "")), "relation": str(r.get("relation", "")), "to": str(r.get("to", ""))}
                for r in body.get("relations", []) if r.get("from") and r.get("to")
            )
            claims.extend(body.get("claims", []))
        if not ie_ok:
            # Degraded but never dead: the deterministic extractor keeps ingestion working and the
            # provenance tag makes the degradation visible instead of silent.
            extractor = "lattice-studio/deterministic-v0-fallback"
            ents, rels = extract_facts(text)
            for n in ents:
                mentions.setdefault(n, "")
            relations.extend({"from": a, "relation": "co_occurs", "to": b} for a, b in rels)

        # 2) RESOLVE — every unique mention through ER; golden records give the canonical name per entity.
        mention_names = list(mentions)
        canonical_of: dict[str, str] = {n: n for n in mention_names}    # identity fallback
        aliases_of: dict[str, list[str]] = {}
        er_meta: dict[str, Any] = {"resolved": False, "merged": 0, "review_queue": 0}
        if mention_names:
            records = [
                {"id": f"m{i}", "name": n, "attributes": ({"type": mentions[n]} if mentions[n] else {}), "scope": coll}
                for i, n in enumerate(mention_names)
            ]
            er_body, er_err = await _req(client, "POST", f"{ER_URL}/resolve", json={"records": records})
            if er_err or not isinstance(er_body, dict):
                if er_err:
                    errors.append(f"entity-resolution: {er_err} (identity fallback)")
            else:
                by_rid = {r["id"]: r["name"] for r in records}
                for ent in er_body.get("entities", []):
                    canon = str(ent.get("canonical", {}).get("name", "")).strip()
                    members = [by_rid[m] for m in ent.get("members", []) if m in by_rid]
                    if not canon or not members:
                        continue
                    for m in members:
                        canonical_of[m] = canon
                    if len(members) > 1:
                        aliases_of[canon] = sorted(m for m in members if m != canon)
                er_meta = {"resolved": True, "merged": er_body.get("merged", 0),
                           "review_queue": len(er_body.get("review_queue", [])),
                           "replay_key": er_body.get("replay_key")}

        # 3) WRITE — canonical nodes, then relations mapped through the canonical ids. A relation is written
        # only when BOTH endpoints resolved to a known entity — dependency spans that matched nothing would
        # otherwise mint junk nodes.
        prov = {"epistemic_mode": "observed", "source": src, "extractor": extractor,
                "project": coll, "kko_type": "Particulars", "doc_sha": doc_sha,
                **({"filename": filename} if filename else {}), **(extra_prov or {})}
        canon_type: dict[str, str] = {}
        for m, canon in canonical_of.items():
            if mentions.get(m) and not canon_type.get(canon):
                canon_type[canon] = mentions[m]
        written_nodes = written_edges = 0
        node_calls = [
            _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                 json={"id": _ent_id(coll, canon),
                       "labels": [coll, "Entity", *([canon_type[canon]] if canon_type.get(canon) else [])],
                       "properties": {"name": canon, **({"aliases": aliases_of[canon]} if canon in aliases_of else {}), **prov}})
            for canon in sorted(set(canonical_of.values()))
        ]
        for _, err in await asyncio.gather(*node_calls):
            if err:
                errors.append(err)
            else:
                written_nodes += 1
        lookup = {n.lower(): c for n, c in canonical_of.items()}
        resolved_rels = []
        for r in relations:
            fa, ta = lookup.get(r["from"].strip().lower()), lookup.get(r["to"].strip().lower())
            if fa and ta and fa != ta:
                label = re.sub(r"[^a-z0-9]+", "_", r["relation"].lower()).strip("_") or "relates_to"
                resolved_rels.append((fa, label, ta))
        skipped_relations = len(relations) - len(resolved_rels)
        edge_calls = [
            _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                 json={"label": label, "from": _ent_id(coll, fa), "to": _ent_id(coll, ta), "properties": prov})
            for fa, label, ta in dict.fromkeys(resolved_rels)
        ]
        for _, err in await asyncio.gather(*edge_calls):
            if err:
                errors.append(err)
            else:
                written_edges += 1

    return {
        "project": project, "projectCollection": coll, "source": src, "doc_sha": doc_sha,
        "chunks": len(chunks),
        "extracted": {"mentions": len(mentions), "relations": len(relations), "claims": len(claims)},
        "resolution": {**er_meta, "entities": len(set(canonical_of.values()))},
        "written": {"nodes": written_nodes, "edges": written_edges, "skipped_relations": skipped_relations},
        "claims": claims[:20],
        "provenance": prov,
        "sample": sorted(set(canonical_of.values()))[:10],
        "errors": errors[:5] or None,
    }


# ---------------------------------------------------------------------------
# File front door: PDF/DOCX/plain-text → text → the same IE→ER→graph pipeline.
# Payload is base64 JSON (matching the estate's JSON-body convention) rather than
# multipart, so callers — Noetica, app-vue, curl — need no special client handling.
MAX_FILE_BYTES = int(os.getenv("STUDIO_MAX_FILE_BYTES", str(20 * 1024 * 1024)))


def _doc_to_text(filename: str, data: bytes) -> str:
    """Convert a supported document to plain text. Raises 415 for unsupported types; parser
    failures are the caller's to wrap (422) so corrupt files never 500."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        from pypdf import PdfReader  # lazy: keeps cold-start light for the non-file routes
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == "docx":
        from docx import Document as DocxDocument  # lazy, same reason
        d = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        # Tables carry the numbers in business documents — flatten row-wise, cells pipe-joined.
        for t in d.tables:
            parts.extend(" | ".join(c.text.strip() for c in row.cells) for row in t.rows)
        return "\n\n".join(x for x in parts if x.strip())
    if ext in ("txt", "md", "markdown"):
        return data.decode("utf-8", errors="replace")
    raise HTTPException(
        status_code=415,
        detail=f"unsupported file type .{ext or '?'} — supported: pdf, docx, txt, md "
               "(for tabular csv/json use /api/studio/ingest)")


class IngestFileRequest(BaseModel):
    project: str = "default"
    filename: str
    content_b64: str
    source: str | None = None


@app.post("/api/studio/ingest-file")
async def ingest_file(req: IngestFileRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """Drop a FILE in, get linked knowledge out — the literal ST024 outcome. Decodes + converts the
    document, then runs the exact same governed pipeline as /ingest-document. Provenance additionally
    carries file_sha (hash of the original bytes) next to doc_sha (hash of the extracted text), so a
    fact is traceable both to the file that was dropped and to the text the extractor actually saw."""
    _require_write_token(authorization)
    try:
        data = base64.b64decode(req.content_b64, validate=True)
    except Exception:
        raise HTTPException(status_code=400, detail="content_b64 is not valid base64") from None
    if not data:
        raise HTTPException(status_code=400, detail="empty file")
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail=f"file exceeds {MAX_FILE_BYTES} bytes")
    try:
        text = _doc_to_text(req.filename, data)
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001 — corrupt/unparseable document must 422, never 500
        raise HTTPException(status_code=422, detail=f"could not parse {req.filename}: {exc}") from None
    if not text.strip():
        raise HTTPException(status_code=422,
                            detail=f"no extractable text in {req.filename} (scanned/image-only PDFs need OCR first)")
    return await _run_ingest_pipeline(
        req.project, text.strip(), req.filename, req.source,
        extra_prov={"file_sha": hashlib.sha256(data).hexdigest()})


# ---------------------------------------------------------------------------
# Fact-mode extraction (IFM stage 03): parse blocks → typed facts against a target
# SQL schema. This serves the contract the compute-gateway extraction adapter speaks
# ({blocks, target_schema} → {facts[]}) — deterministic, span-preserving, and it
# NEVER fabricates: a field the document doesn't state yields no fact at all.
# left boundary blocks period/word-glued digits ('FY26', 'v1.5') from reading as values
_FACT_NUM = re.compile(
    r"(?<![\w.$])\(\s*\$?\s*\d[\d,]*(?:\.\d+)?\s*(?:%|bn|b|m|k)?\s*\)"     # accounting negative: (…)
    r"|(?<![\w.$])\$?\s*-?\d[\d,]*(?:\.\d+)?\s*(?:%|bn|b|m|k)?(?=[\s|)]|$)", re.I)
_FACT_SUFFIX_UNIT = {"%": "%", "m": "m", "b": "bn", "bn": "bn", "k": "k"}


def _parse_fact_value(cell: str) -> tuple[float, str | None] | None:
    """First numeric token in a cell → (value, detected_unit|None). Handles $, thousands
    commas, m/bn/k magnitude suffixes, %, and accounting-style (parenthesised) negatives.
    Values stay AS PRINTED (1204 for '$1,204m', unit 'm') — normalisation is validate's job."""
    m = _FACT_NUM.search(cell)
    if not m:
        return None
    tok = m.group(0).strip()
    neg = tok.startswith("(")
    t = tok.strip("()").replace("$", "").replace(",", "").strip()
    m2 = re.match(r"(-?\d+(?:\.\d+)?)\s*(%|bn|b|m|k)?$", t, re.I)
    if not m2:
        return None
    val = float(m2.group(1))
    if neg:
        val = -val
    sfx = (m2.group(2) or "").lower()
    return val, (_FACT_SUFFIX_UNIT.get(sfx) if sfx else None)


def _fact_label_variants(field: dict[str, Any]) -> list[str]:
    """Label strings that count as this field, all normalised: schema-author aliases
    first (explicit intent wins), then the field name with underscores as spaces."""
    out = [str(v).strip().lower() for v in field.get("labels", []) if str(v).strip()]
    out.append(str(field.get("name", "")).replace("_", " ").strip().lower())
    return [v for v in out if v]


def _norm_label(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().rstrip(":").lower())


def _pick_period_cell(cells: list[str], header: list[str], period: str | None) -> tuple[float, str | None] | None:
    """Which column is THE value? If the table header names the requested period ('FY26'),
    that column wins. Otherwise the LAST parsable cell — results tables put the current
    period rightmost by convention. (Wrong column = silently trading on last year's number.)"""
    if period and header:
        p = period.strip().lower()
        for i, h in enumerate(header):
            if p and p in h.strip().lower() and i < len(cells):
                if (parsed := _parse_fact_value(cells[i])) is not None:
                    return parsed
    for c in reversed(cells[1:]):
        if (parsed := _parse_fact_value(c)) is not None:
            return parsed
    return None


def extract_facts_from_blocks(blocks: list[dict[str, Any]], target_schema: dict[str, Any],
                              period: str | None = None) -> list[dict[str, Any]]:
    """For each schema field, find the best-evidenced value in the blocks. Table rows beat
    prose (a labelled cell IS the number; prose needs interpretation), exact label matches
    beat 'Total X' prefixes, earlier pages beat later. One fact per field, or none."""
    fields = [f for f in target_schema.get("fields", []) if f.get("name")]
    facts: list[dict[str, Any]] = []
    for field in fields:
        variants = _fact_label_variants(field)
        best: dict[str, Any] | None = None
        for b in blocks:
            page, kind = b.get("page"), b.get("kind", "text")
            header: list[str] = []
            for lno, line in enumerate(str(b.get("text", "")).splitlines()):
                if kind == "table" and "|" in line:
                    cells = [c.strip() for c in line.split("|")]
                    if lno == 0 and _parse_fact_value(line) is None:
                        header = cells          # a numberless first row is the period header
                    lbl = _norm_label(cells[0])
                    conf = 0.95 if lbl in variants else 0.9 if any(lbl == f"total {v}" for v in variants) else None
                    if conf is None:
                        continue
                    parsed = _pick_period_cell(cells, header, period)
                    span = f"p{page}/tbl: {cells[0]}"
                else:
                    low = line.lower()
                    hit = next((v for v in variants if v in low), None)
                    if hit is None:
                        continue
                    # the number must FOLLOW the label in the same line — 'revenue of $500m'
                    parsed = _parse_fact_value(line[low.index(hit) + len(hit):])
                    conf, span = 0.6, f"p{page}/txt: {line.strip()[:80]}"
                if parsed is None:
                    continue
                value, detected_unit = parsed
                cand = {"field": field["name"], "value": value,
                        "unit": detected_unit or field.get("unit"),
                        "page": page, "source_span": span, "confidence": conf, "verbatim": True}
                if best is None or (cand["confidence"], -(cand["page"] or 0)) > (best["confidence"], -(best["page"] or 0)):
                    best = cand
        if best is not None:
            facts.append(best)
    return facts


class ExtractFactsRequest(BaseModel):
    project: str = "default"
    blocks: list[dict[str, Any]]
    target_schema: dict[str, Any] = {}
    period: str | None = None
    document: str | None = None


@app.post("/api/studio/extract-facts")
async def extract_facts_endpoint(req: ExtractFactsRequest) -> dict[str, Any]:
    """The gateway extraction backend. Pure and side-effect-free — writes nothing, calls
    nothing — so it carries no write gate; the compute plane's entitlement gate governs
    who may run it as a compute. Every fact keeps page + source span; absent fields are
    absent, not invented — the downstream warrant machinery depends on that honesty."""
    facts = extract_facts_from_blocks(req.blocks, req.target_schema, req.period)
    return {"facts": facts,
            "fields_requested": len(req.target_schema.get("fields", [])),
            "fields_found": len(facts)}


class NodeRequest(BaseModel):
    project: str = "default"
    name: str
    labels: list[str] | None = None
    epistemic_mode: str = "observed"
    source: str | None = None


class EdgeRequest(BaseModel):
    project: str = "default"
    from_name: str
    to_name: str
    label: str = "relates_to"
    epistemic_mode: str = "observed"
    source: str | None = None


def _workbench_prov(coll: str, mode: str, source: str | None) -> dict[str, Any]:
    """Provenance for a HAND-AUTHORED fact — same shape /extract stamps, so a manual node is as proof-carrying
    as an extracted one. extractor marks it as workbench-authored; epistemic_mode is the user's assertion."""
    return {"epistemic_mode": mode, "source": source or "workbench",
            "extractor": "lattice-studio/workbench-v0", "project": coll, "kko_type": "Particulars"}


@app.post("/api/studio/node")
async def add_node(req: NodeRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """WRITE workbench: add ONE node to the project graph by hand. Same fail-closed token, project scope, and
    provenance as /extract — a hand-authored entity lands in the same graph, carrying its epistemic status."""
    _require_write_token(authorization)
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="name required")
    coll = proj_collection(req.project)
    nid = _ent_id(coll, name)
    prov = _workbench_prov(coll, req.epistemic_mode, req.source)
    # keep the project + Entity labels canonical; append any extra user labels (deduped)
    labels = [coll, "Entity"] + [l for l in (req.labels or []) if l not in (coll, "Entity")]
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        _, err = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                            json={"id": nid, "labels": labels, "properties": {"name": name, **prov}})
    if err:
        raise HTTPException(status_code=502, detail=f"graph write failed: {err}")
    return {"id": nid, "name": name, "project": req.project, "projectCollection": coll,
            "labels": labels, "provenance": prov, "written": True}


@app.post("/api/studio/edge")
async def add_edge(req: EdgeRequest, authorization: str = Header(default="")) -> dict[str, Any]:
    """WRITE workbench: add ONE edge to the project graph by hand. Upserts both endpoints first (idempotent),
    so an edge between new names just works; then writes the relation. Same token/scope/provenance as /extract."""
    _require_write_token(authorization)
    a, b = req.from_name.strip(), req.to_name.strip()
    if not a or not b:
        raise HTTPException(status_code=422, detail="from_name and to_name required")
    coll = proj_collection(req.project)
    fid, tid = _ent_id(coll, a), _ent_id(coll, b)
    label = _norm(req.label).replace(" ", "_") or "relates_to"
    prov = _workbench_prov(coll, req.epistemic_mode, req.source)
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        # endpoints first (idempotent upsert), then the edge — mirrors /extract's nodes-before-edges order.
        for nm, nid in ((a, fid), (b, tid)):
            _, err = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/node",
                                json={"id": nid, "labels": [coll, "Entity"], "properties": {"name": nm, **prov}})
            if err:
                raise HTTPException(status_code=502, detail=f"graph write failed (node {nm}): {err}")
        _, err = await _req(client, "POST", f"{HELLGRAPH_URL}/api/graph/edge",
                            json={"label": label, "from": fid, "to": tid, "properties": prov})
    if err:
        raise HTTPException(status_code=502, detail=f"graph write failed (edge): {err}")
    return {"from": fid, "to": tid, "label": label, "project": req.project,
            "projectCollection": coll, "provenance": prov, "written": True}


def _map_node(n: Any) -> dict[str, Any]:
    props = n.get("properties", n) if isinstance(n, dict) else {}
    nid = n.get("id", "") if isinstance(n, dict) else str(n)
    return {
        "id": nid, "name": props.get("name", nid),
        "epistemic_mode": props.get("epistemic_mode", "unknown"),
        "source": props.get("source"), "extractor": props.get("extractor"),
        "kko_type": props.get("kko_type", "Particulars"),  # KKO upper-ontology type
        "labels": n.get("labels", []) if isinstance(n, dict) else [],
    }


async def _fetch_subgraph(coll: str, limit: int = 200) -> tuple[list[dict[str, Any]], list[dict[str, Any]], str | None]:
    """Read the project's INDUCED SUBGRAPH (nodes + internal edges) from the live HellGraph kernel.

    The kernel returns an induced subgraph (edges only where both endpoints are in the node set), so the
    explorer draws real topology — not a chip-cloud — and every node still carries its provenance.
    """
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        res, err = await _req(client, "GET", f"{HELLGRAPH_URL}/api/graph/subgraph?label={coll}&limit={limit}")
    raw_nodes = res.get("nodes", []) if isinstance(res, dict) else []
    raw_edges = res.get("edgeList", []) if isinstance(res, dict) else []
    nodes = [_map_node(n) for n in (raw_nodes if isinstance(raw_nodes, list) else [])[:limit]]
    edges = [
        {"id": e.get("id", ""), "source": e.get("from", ""), "target": e.get("to", ""),
         "label": e.get("label", ""), "weight": (e.get("properties") or {}).get("n", 1)}
        for e in (raw_edges if isinstance(raw_edges, list) else [])
        if isinstance(e, dict) and e.get("from") and e.get("to")
    ]
    return nodes, edges, err


async def _fetch_nodes(coll: str, limit: int = 200) -> tuple[list[dict[str, Any]], str | None]:
    """Node-only read (RDF export path) — delegates to the subgraph fetch and drops edges."""
    nodes, _edges, err = await _fetch_subgraph(coll, limit)
    return nodes, err


@app.get("/api/studio/graph")
async def graph(project: str = "default", limit: int = 100, _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """KE-2: the project sub-graph with PROVENANCE PER NODE — the differentiator, read from the live kernel.

    Not just the node, but its epistemic status + source + extractor, in one governed project scope — what a
    Neo4j/Bloom explorer can't show natively.
    """
    coll = proj_collection(project)
    nodes, edges, err = await _fetch_subgraph(coll, limit)
    dist: dict[str, int] = {}
    for x in nodes:
        dist[x["epistemic_mode"]] = dist.get(x["epistemic_mode"], 0) + 1
    return {"project": project, "projectCollection": coll, "nodes": nodes, "edges": edges,
            "count": len(nodes), "edge_count": len(edges),
            "epistemic_distribution": dist, "degraded": (err if err else None)}


def _derivation_summary(node: dict[str, Any], derivations: list[dict[str, Any]]) -> str:
    """One human sentence: how this fact came to be known — its epistemic status, who/what produced it, and how
    many other facts it's connected to. The 'How derived?' answer a Bloom/Stardog node inspector can't give."""
    mode = node.get("epistemic_mode", "unknown")
    by = node.get("extractor") or "an unknown process"
    src = node.get("source")
    origin = f" from {src}" if src else ""
    n = len(derivations)
    rel = "no other facts yet" if n == 0 else f"{n} related fact{'s' if n != 1 else ''}"
    return f"‘{node.get('name', node['id'])}’ is held as {mode}, produced by {by}{origin}, and connected to {rel}."


@app.get("/api/studio/provenance")
async def provenance(project: str = "default", id: str = "", _auth: dict[str, Any] | None = Depends(require_read)) -> dict[str, Any]:
    """KE-5: 'How derived?' — the derivation of ONE fact. Its provenance (epistemic status + source + extractor +
    KKO upper-ontology type) plus the edges that connect it (what it was co-observed / reasoned with), read from
    the live project subgraph. This is the proof-carrying lineage a Neo4j Bloom / Stardog node inspector can't
    show: not just properties, but where the fact came from and how well it's known."""
    if not id:
        raise HTTPException(status_code=422, detail="id required")
    coll = proj_collection(project)
    nodes, edges, err = await _fetch_subgraph(coll, 500)
    node = next((n for n in nodes if n["id"] == id), None)
    if node is None:
        return {"id": id, "project": project, "projectCollection": coll, "found": False, "degraded": err}
    by_id = {n["id"]: n for n in nodes}
    derivations = []
    for e in edges:
        if e["source"] != id and e["target"] != id:
            continue
        out = e["source"] == id
        other = by_id.get(e["target"] if out else e["source"], {})
        derivations.append({
            "relation": e["label"], "direction": "out" if out else "in",
            "with": {"id": other.get("id", ""), "name": other.get("name", ""),
                     "epistemic_mode": other.get("epistemic_mode", "unknown"), "source": other.get("source")},
            "weight": e.get("weight", 1),
        })
    derivations.sort(key=lambda d: d["weight"], reverse=True)
    return {
        "id": id, "project": project, "projectCollection": coll, "found": True, "name": node["name"],
        "epistemic_mode": node["epistemic_mode"], "source": node.get("source"),
        "extractor": node.get("extractor"), "kko_type": node.get("kko_type", "Particulars"),
        "labels": node.get("labels", []),
        "derivations": derivations, "derivation_count": len(derivations),
        "summary": _derivation_summary(node, derivations),
        "degraded": err,
    }


@app.get("/api/studio/graph.ttl")
async def graph_ttl(project: str = "default", limit: int = 500, _auth: dict[str, Any] | None = Depends(require_read)) -> Response:
    """KE-3: RDF/Turtle export — standards interop (Protégé / GraphDB / Anzo / Stardog) that CARRIES provenance.

    Every node exports as a PROV-O-annotated resource: rdf:type sp:Entity, rdfs:label, sp:epistemicMode,
    dct:source, prov:wasGeneratedBy. Their RDF exports drop provenance; ours doesn't — epistemic status + origin
    ride the standard triples, so a proof-carrying graph stays proof-carrying when it leaves us.
    """
    from rdflib import Graph as RDFGraph, Literal, Namespace, URIRef
    from rdflib.namespace import DCTERMS, PROV, RDF, RDFS

    coll = proj_collection(project)
    nodes, edges, _ = await _fetch_subgraph(coll, limit)   # nodes AND edges — the reasoner needs relations
    g = RDFGraph()
    # KKO (KBpedia Knowledge Ontology) is the estate's upper ontology — the export types INTO it, so the graph
    # is grounded in a formal, open (CC-BY-4.0), Peircean upper ontology that Protégé/GraphDB/Anzo/Stardog already
    # understand. Meet them on standards (KKO/RDF/PROV) AND carry the provenance moat they drop on export.
    KKO = Namespace("http://kbpedia.org/ontologies/kko#")
    SP = Namespace("https://socioprophet.ai/kg#")
    PROJ = Namespace(f"https://socioprophet.ai/kg/{coll}/")
    g.bind("kko", KKO); g.bind("sp", SP); g.bind("proj", PROJ); g.bind("prov", PROV); g.bind("dct", DCTERMS)

    def _local(node_id: str) -> URIRef:
        return PROJ[(node_id.split(":")[-1] or "node")]

    for n in nodes:
        u = _local(n["id"])
        g.add((u, RDF.type, KKO[n.get("kko_type", "Particulars")]))  # KKO upper-ontology type (Peircean)
        g.add((u, RDFS.label, Literal(n["name"])))
        # epistemic_mode is a Peircean inference status (KKO kko:Methodeutic) — provenance grounded in the standard.
        g.add((u, SP.epistemicMode, Literal(n["epistemic_mode"])))
        if n.get("source"): g.add((u, DCTERMS.source, Literal(n["source"])))
        # PROV-O: wasGeneratedBy must reference a prov:Activity RESOURCE, not a bare literal.
        if n.get("extractor"):
            act = SP[f"activity/{str(n['extractor']).replace(' ', '_')}"]
            g.add((act, RDF.type, PROV.Activity))
            g.add((u, PROV.wasGeneratedBy, act))
    # Edges — WITHOUT these the exported graph is a bag of typed nodes with NO relations, and a reasoner
    # pulling graph.ttl has nothing to reason over. Each edge becomes a real predicate triple; the label
    # (e.g. co_occurs) is minted under the sp: vocabulary so it is a dereferenceable property.
    for e in edges:
        src, tgt, label = e.get("source"), e.get("target"), e.get("label") or "relatedTo"
        if not src or not tgt:
            continue
        g.add((_local(src), SP[str(label)], _local(tgt)))
    return Response(content=g.serialize(format="turtle"), media_type="text/turtle")


# ── Notebook runtime proxy → lattice-forge (isolated sovereign-runtime) ──────
# The BFF is the ONLY thing that talks to the forge; forge/JupyterLab are never
# exposed directly. Governance (receipts) lives in the forge; we relay + surface.
def _forge_headers() -> dict[str, str]:
    return {"Authorization": f"Bearer {FORGE_TOKEN}"} if FORGE_TOKEN else {}


async def _forge(method: str, path: str, *, json: Any | None = None, params: dict | None = None) -> tuple[Any, str | None]:
    try:
        async with httpx.AsyncClient(timeout=FORGE_TIMEOUT, headers=_forge_headers()) as c:
            r = await c.request(method, f"{FORGE_URL}{path}", json=json, params=params)
            if r.status_code == 200:
                return r.json(), None
            return None, f"forge HTTP {r.status_code}"
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


class NotebookSession(BaseModel):
    project: str = "default"
    adapter: str | None = None
    name: str | None = None


class NotebookExec(BaseModel):
    project: str = "default"
    code: str
    language: str = "python"
    adapter: str | None = None
    session_id: str | None = None


@app.get("/api/studio/notebook/adapters")
async def nb_adapters() -> dict[str, Any]:
    data, err = await _forge("GET", "/v1/adapters")
    return data or {"degraded": err, "default": "jupyterlab", "adapters": {}}


@app.post("/api/studio/notebook/session")
async def nb_create_session(req: NotebookSession) -> dict[str, Any]:
    data, err = await _forge("POST", "/v1/session",
                             json={"project": req.project, "adapter": req.adapter, "name": req.name})
    if err:
        return {"degraded": err}
    return data


@app.get("/api/studio/notebook/sessions")
async def nb_sessions(project: str = "default") -> dict[str, Any]:
    data, err = await _forge("GET", "/v1/sessions", params={"project": project})
    return data or {"project": project, "sessions": [], "degraded": err}


@app.post("/api/studio/notebook/execute")
async def nb_execute(req: NotebookExec) -> dict[str, Any]:
    data, err = await _forge("POST", "/v1/execute", json=req.model_dump())
    if err:
        # honest degradation — the surface shows "runtime unavailable", never a fake result
        return {"status": "degraded", "degraded": err, "outputs": [], "receipt": None}
    return data


@app.get("/api/studio/notebook/receipts")
async def nb_receipts(project: str = "default") -> dict[str, Any]:
    data, err = await _forge("GET", "/v1/receipts", params={"project": project})
    return data or {"project": project, "receipts": [], "count": 0, "degraded": err}


# ── Universal Compute Plane proxy → compute-gateway ──────────────────────────
# Any kind of compute (notebook, graph, spark, inference) through ONE governed
# door: gate → route → seal a signed receipt → type the warrant → write provenance.
# The BFF is the only thing that talks to the gateway; the gateway routes to the
# isolated backends. The compute plane's story, surfaced.
def _gateway_headers() -> dict[str, str]:
    return {"Authorization": f"Bearer {GATEWAY_TOKEN}"} if GATEWAY_TOKEN else {}


async def _gateway(method: str, path: str, *, json: Any | None = None, params: dict | None = None) -> tuple[Any, str | None]:
    try:
        async with httpx.AsyncClient(timeout=GATEWAY_TIMEOUT, headers=_gateway_headers()) as c:
            r = await c.request(method, f"{COMPUTE_GATEWAY_URL}{path}", json=json, params=params)
            if r.status_code == 200:
                return r.json(), None
            return None, f"gateway HTTP {r.status_code}"
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


class ComputeRun(BaseModel):
    kind: str
    spec: dict[str, Any] = {}
    project: str = "default"
    backend: str | None = None
    entitlement: str | None = None


@app.get("/api/studio/compute/registry")
async def compute_registry(project: str = "default") -> dict[str, Any]:
    data, err = await _gateway("GET", "/v1/registry", params={"project": project})
    return data or {"project": project, "kinds": [], "degraded": err}


@app.post("/api/studio/compute/run")
async def compute_run(req: ComputeRun) -> dict[str, Any]:
    data, err = await _gateway("POST", "/v1/compute", json=req.model_dump())
    if err:
        return {"status": "degraded", "degraded": err, "outputs": [], "receipt": None}
    return data


@app.get("/api/studio/compute/receipts")
async def compute_receipts(project: str = "default") -> dict[str, Any]:
    data, err = await _gateway("GET", "/v1/receipts", params={"project": project})
    return data or {"project": project, "receipts": [], "count": 0, "degraded": err}


class ComputePlan(BaseModel):
    capabilities: list[str] = []
    project: str = "default"
    intent: str | None = None
    entitlement: str | None = None


@app.post("/api/studio/compute/plan")
async def compute_plan(req: ComputePlan) -> dict[str, Any]:
    """Plan a governed workflow over the capability registry (the agent action space).
    A preview — returns a runnable workflow spec the surface hands back to /run."""
    data, err = await _gateway("POST", "/v1/plan", json=req.model_dump())
    return data or {"strategy": "degraded", "degraded": err, "plan": None, "steps": []}
